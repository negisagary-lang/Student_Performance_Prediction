"""
Week 4 - Comprehensive Data Science Report and Insights Presentation Plan
DOCX Report Generator
Primary deliverable: Week_4_Comprehensive_Data_Science_Report.docx

Honesty note: This is a planning/communication report. No real dataset was
processed and no model was trained. All findings, charts and numbers are
HYPOTHETICAL / ILLUSTRATIVE and clearly labelled as such.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = r"C:\Users\Lenovo\Student_Performance_Prediction\Week4_Report"
ASSETS = os.path.join(OUTPUT_DIR, "assets")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Week_4_Comprehensive_Data_Science_Report.docx")

# Personal details (verified from Week 1/3 - Sagar Negi internship)
STUDENT_NAME = "Sagar Negi"
ROLL_NUMBER = "[Roll Number]"
COLLEGE = "Faculty of Technology, Veer Madho Singh Bhandari"
UNIVERSITY = "Uttarakhand Technical University (UTU)"
DEPARTMENT = "[Department]"
ORG = "Yuva Internship"
PROGRAM = "Virtual Data Science Explorer Intern"
MENTOR = "[Mentor / Supervisor]"
SUBMISSION_DATE = "31 August 2026"

doc = Document()

# ---- A4 page setup ----
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ---- Header / footer ----
section.different_first_page_header_footer = True
hp = section.header.paragraphs[0]
hp.text = "Week 4 | Comprehensive Data Science Report and Insights Presentation Plan"
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in hp.runs:
    run.font.size = Pt(9); run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

fp = section.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
run.font.size = Pt(9); run.font.name = 'Calibri'
run.italic = True
run.add_text(" Student Performance Analysis & Prediction ")
fld = parse_xml(r'<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>')
fp._p.append(fld)

# ---- Styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for level, size, color in [(1, 17, '1A3C6E'), (2, 14, '2C5F8A'), (3, 12, '344E6B')]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'; hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color); hs.font.bold = True

def heading(text, level):
    doc.add_heading(text, level=level)

def para(text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.name = 'Calibri'
    run.bold = bold; run.italic = italic
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.15
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ' '); r.font.size = Pt(11)
        r.font.name = 'Calibri'; r.bold = True
    r = p.add_run(text); r.font.size = Pt(11); r.font.name = 'Calibri'
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = p.add_run(bold_prefix + ' '); r.font.size = Pt(11)
        r.font.name = 'Calibri'; r.bold = True
    r = p.add_run(text); r.font.size = Pt(11); r.font.name = 'Calibri'
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = 'Consolas'; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>'))
    return p

def set_cell_shading(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def table(headers, rows, widths=None, caption_text=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9.5); run.font.name = 'Calibri'
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A3C6E')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]; cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9); run.font.name = 'Calibri'
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'EBF1F8')
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    if caption_text:
        caption(caption_text)
    return t

def caption(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.font.size = Pt(10); run.font.name = 'Calibri'
    run.italic = True; run.bold = True
    p.paragraph_format.space_after = Pt(10)
    return p

def figure(path, caption_text, width=6.3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(path, width=Inches(width))
    caption(caption_text)

def callout(text, color='EEF4FB', border='1A3C6E'):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.size = Pt(10.5); run.font.name = 'Calibri'; run.italic = True
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="4" w:color="{border}"/></w:pBdr>'))
    p.paragraph_format.space_after = Pt(8)
    return p

def centered(text, size, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.font.size = Pt(size); run.font.name = 'Calibri'
    run.bold = bold; run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p

def page_break():
    doc.add_page_break()

# =================================================================
# COVER PAGE
# =================================================================
for _ in range(4):
    doc.add_paragraph()

centered('WEEK 4 INTERNSHIP REPORT', 14, bold=True, space_after=20)
centered('Student Performance Analysis & Prediction', 24, bold=True, space_after=8)
centered('Week 4 Internship Report', 16, bold=True, space_after=6)
centered('Comprehensive Data Science Report and Insights Presentation Plan', 13, bold=True, space_after=16)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 55); run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E); run.font.size = Pt(10)

doc.add_paragraph()
cover = [
    ('Student Name', STUDENT_NAME),
    ('Roll Number', ROLL_NUMBER),
    ('College / Faculty', COLLEGE),
    ('University', UNIVERSITY),
    ('Department', DEPARTMENT),
    ('Internship Organization', ORG),
    ('Internship Program', PROGRAM),
    ('Mentor / Supervisor', MENTOR),
    ('Submission Date', SUBMISSION_DATE),
]
for label, value in cover:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ':  '); r1.font.size = Pt(12); r1.font.name = 'Calibri'; r1.bold = True
    r2 = p.add_run(value); r2.font.size = Pt(12); r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph(); doc.add_paragraph()
centered('This report is a planning and communication exercise. The dataset and all findings are', 10, italic=True)
centered('HYPOTHETICAL / ILLUSTRATIVE and are labelled as such throughout.', 10, italic=True)

page_break()

# =================================================================
# TABLE OF CONTENTS
# =================================================================
heading('Table of Contents', 1)
toc_p = doc.add_paragraph()
r1 = toc_p.add_run()._r
r1.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
r2 = toc_p.add_run()._r
r2.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
r3 = toc_p.add_run()._r
r3.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
r4 = toc_p.add_run()._r
r4.append(parse_xml('<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">Right-click here and choose "Update Field" to build the Table of Contents.</w:t>'))
r5 = toc_p.add_run()._r
r5.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
page_break()

# =================================================================
# 1. EXECUTIVE SUMMARY
# =================================================================
heading('1. Executive Summary', 1)
para(
    'The Student Performance Analysis and Prediction project applies Data Science and Machine '
    'Learning to understand the factors associated with student academic performance and to '
    'develop a framework for predicting academic outcomes. Educational institutions frequently '
    'hold large amounts of student information, yet they often lack systematic, data-driven methods '
    'to convert that information into reliable insights and actionable support.'
)
para(
    'This Week 4 report consolidates the complete proposed Data Science workflow into one '
    'professional, executive-level communication document. It brings together the problem '
    'definition, proposed data, methodology, exploratory analysis approach, and the machine-learning '
    'framework designed in earlier phases, and translates them into the language of project managers, '
    'school administrators, academic coordinators, and educational stakeholders.'
)
para(
    'The data that would be analysed includes proposed variables such as study hours, attendance, '
    'previous scores, assignment performance, absences, age, internet access, extracurricular '
    'activity, educational background, and the final score. A combination of classical statistics, '
    'exploratory data analysis, visualization, and predictive modelling would be used to surface '
    'patterns, generate insights, and support forecasting of academic outcomes.'
)
para(
    'Because no real dataset has been provided for this exercise, every finding and visual in this '
    'report is deliberately hypothetical and clearly identified as such. The mock-up charts '
    'demonstrate how genuine insights would be packaged once real data becomes available. The '
    'executive structure, storytelling framework, and presentation strategy in this report can be '
    'used directly to communicate the completed project to non-technical stakeholders.'
)
callout(
    'This report presents a proposed analytical and communication framework using illustrative '
    'findings to demonstrate how the completed project could be presented once real data becomes '
    'available. No analysis was executed and no real results are claimed.'
)

page_break()

# =================================================================
# 2. PROJECT OVERVIEW
# =================================================================
heading('2. Project Overview', 1)
para(
    'The overall project is Student Performance Analysis & Prediction. Its purpose is to analyse '
    'the factors associated with student academic performance and to develop a framework for '
    'predicting academic outcomes using Data Science and Machine Learning techniques in Python.'
)
para('The following variables are proposed for the project (illustrative only):')
table(
    ['Variable', 'Proposed Role'],
    [
        ['Study Hours', 'Feature'],
        ['Attendance', 'Feature'],
        ['Previous Score', 'Feature'],
        ['Assignment Performance', 'Feature'],
        ['Absences', 'Feature'],
        ['Age', 'Feature'],
        ['Internet Access', 'Feature'],
        ['Extracurricular Activities', 'Feature'],
        ['Educational Background', 'Feature'],
        ['Final Score', 'Target'],
    ],
    widths=[3.6, 1.6],
    caption_text='Table 1: Proposed variables for the project (illustrative)',
)
callout('The list above is proposed and illustrative. The final set of variables depends on the actual dataset.')

# =================================================================
# 3. BUSINESS / EDUCATIONAL PROBLEM
# =================================================================
heading('3. Business / Educational Problem', 1)
para(
    'From a stakeholder perspective, the central challenge is this: educational institutions may '
    'collect large volumes of student information but lack a systematic, scalable way to identify '
    'the patterns associated with academic performance. Raw data alone does not tell a clear story.'
)
para('Data Science could help educational stakeholders to:')
bullet('Identify patterns in student data.')
bullet('Understand factors potentially associated with performance.')
bullet('Detect potential academic risks early.')
bullet('Support data-driven decision making.')
bullet('Improve academic planning.')
bullet('Support targeted interventions.')
callout(
    'Important note on language: this project does not claim that any particular factor CAUSES '
    'poor performance. Findings are described as "may be associated with" academic outcomes. '
    'Correlation is not treated as causation.'
)

# =================================================================
# 4. PROJECT OBJECTIVES
# =================================================================
heading('4. Project Objectives', 1)
heading('4.1 Primary Objective', 2)
para('Develop a Data Science framework for analysing and predicting student academic performance.')
heading('4.2 Secondary Objectives', 2)
numbered('Understand the proposed student-performance data.')
numbered('Identify the important variables.')
numbered('Explore distributions and relationships.')
numbered('Detect data-quality issues.')
numbered('Develop predictive models.')
numbered('Compare candidate models.')
numbered('Evaluate model performance.')
numbered('Generate understandable insights.')
numbered('Communicate results visually.')
numbered('Provide actionable recommendations.')

page_break()

# =================================================================
# 5. METHODOLOGY OVERVIEW
# =================================================================
heading('5. Methodology Overview', 1)
para(
    'This project follows the standard, integrated Data Science lifecycle. Rather than being '
    'treated as separate weekly tasks, the stages below form one continuous, repeatable process '
    'for turning raw educational data into trustworthy insights and responsible actions.'
)
code(
    'Problem Definition\n'
    '       \u2193\n'
    'Data Collection\n'
    '       \u2193\n'
    'Data Understanding\n'
    '       \u2193\n'
    'Data Cleaning\n'
    '       \u2193\n'
    'Exploratory Data Analysis\n'
    '       \u2193\n'
    'Feature Engineering\n'
    '       \u2193\n'
    'Feature Selection\n'
    '       \u2193\n'
    'Machine Learning\n'
    '       \u2193\n'
    'Model Evaluation\n'
    '       \u2193\n'
    'Insight Generation\n'
    '       \u2193\n'
    'Visualization\n'
    '       \u2193\n'
    'Recommendations\n'
    '       \u2193\n'
    'Communication'
)
caption('Figure 1: Integrated Data Science methodology')
for label, desc in [
    ('Problem Definition', 'Clarifying the question that the analysis must answer.'),
    ('Data Collection', 'Gathering the information needed to investigate the problem.'),
    ('Data Understanding', 'Reviewing structure, types, ranges and quality of the data.'),
    ('Data Cleaning', 'Correcting missing, duplicate or invalid information.'),
    ('EDA', 'Exploring patterns, distributions and relationships visually and statistically.'),
    ('Feature Engineering', 'Creating new informative variables from raw data.'),
    ('Feature Selection', 'Retaining informative variables and removing noise.'),
    ('Machine Learning', 'Training candidate predictive models on the data.'),
    ('Model Evaluation', 'Measuring how well models generalise to unseen data.'),
    ('Insight Generation', 'Translating model and analysis output into meaningful conclusions.'),
    ('Visualization', 'Presenting findings clearly for diverse audiences.'),
    ('Recommendations', 'Turning insights into practical, stakeholder-facing actions.'),
    ('Communication', 'Sharing the story in plain, accessible language.'),
]:
    para(desc, bold=False)

page_break()

# =================================================================
# 6. DATA UNDERSTANDING
# =================================================================
heading('6. Data Understanding', 1)
para(
    'This section describes the proposed structure of the dataset. The table below is an '
    'illustrative data dictionary and not based on any real file.'
)
table(
    ['Variable', 'Data Type', 'Role', 'Example Purpose'],
    [
        ['Study Hours', 'Numerical', 'Feature', 'Academic effort'],
        ['Attendance', 'Numerical', 'Feature', 'Participation'],
        ['Previous Score', 'Numerical', 'Feature', 'Prior achievement'],
        ['Absences', 'Numerical', 'Feature', 'Attendance behavior'],
        ['Age', 'Numerical', 'Feature', 'Demographic context'],
        ['Internet Access', 'Categorical', 'Feature', 'Resource availability'],
        ['Final Score', 'Numerical', 'Target', 'Performance outcome'],
    ],
    widths=[1.9, 1.1, 0.9, 2.1],
    caption_text='Table 2: Illustrative data dictionary (proposed)',
)
para(
    'Understanding the data means reviewing each variable for its type (numerical or categorical), '
    'its intended role (feature or target), the values it can take, and the quality of those values. '
    'This understanding informs the cleaning and analysis that follow.'
)

# =================================================================
# 7. DATA PREPARATION
# =================================================================
heading('7. Data Preparation', 1)
para('Before analysis, raw data must be prepared. Each step below is proposed for this project.')
bullet('Missing values: handle absent values appropriately.')
bullet('Duplicate records: remove repeated rows.')
bullet('Invalid values: correct or remove impossible entries.')
bullet('Outliers: examine extreme values for validity.')
bullet('Data-type correction: ensure each column has the right type.')
bullet('Categorical encoding: convert categories to usable form.')
bullet('Feature scaling: standardize numerical range where needed.')
bullet('Feature engineering: create informative new variables.')
bullet('Feature selection: keep only useful variables.')
bullet('Data leakage prevention: keep test/training information separated.')
para(
    'Each step matters because the quality of the insights depends directly on the quality of the '
    'data. Poor cleaning, inconsistent types, or information leakage can quietly distort every '
    'later finding and prediction.'
)

page_break()

# =================================================================
# 8. EXPLORATORY ANALYSIS
# =================================================================
heading('8. Exploratory Analysis', 1)
para(
    'Exploratory Data Analysis (EDA) is used to inspect the structure of the data, detect '
    'quality problems, and build an initial picture of the relationships within it.'
)
heading('8.1 Univariate Analysis', 2)
para('Studies individual variables in isolation, examining their distributions, ranges and central tendencies.')
heading('8.2 Bivariate Analysis', 2)
para('Studies the relationship between two variables at a time, for example attendance and final score.')
heading('8.3 Multivariate Analysis', 2)
para('Studies relationships that involve multiple variables simultaneously.')
heading('8.4 Correlation Analysis', 2)
para(
    'Measures the statistical association between numerical variables. Correlation is useful for '
    'spotting potential relationships, but it must never be interpreted as proof of causation. '
    'No real correlation values are reported in this planning document.'
)
callout(
    'The exploratory analysis described here is proposed. No real correlations or statistics are '
    'claimed because no dataset has been analysed.'
)

# =================================================================
# 9. MACHINE LEARNING APPROACH
# =================================================================
heading('9. Machine Learning Approach', 1)
para(
    'The proposed predictive task is student performance prediction framed as a regression problem, '
    'because the final score is a continuous numerical value.'
)
para('Potential candidate algorithms:')
bullet('Linear Regression - a simple, interpretable baseline.')
bullet('Decision Tree Regression - captures some nonlinear relationships.')
bullet('Random Forest Regression - a strong general-purpose candidate.')
bullet('Gradient Boosting Regression - a powerful advanced candidate.')
para(
    'Multiple models should be compared rather than committing to one early. A baseline model '
    'is necessary so that more complex models can be judged against a fair reference. Random '
    'Forest may be a strong candidate because student performance often depends on nonlinear '
    'interactions between variables; however, the final choice must be based on empirical '
    'evaluation on real data.'
)
callout('No model was trained and no model achieved any claimed score. Model selection is a future, data-driven step.')

# =================================================================
# 10. MODEL EVALUATION
# =================================================================
heading('10. Model Evaluation', 1)
para('For a regression problem the following metrics would be used to evaluate models:')
bullet('MAE - Mean Absolute Error.')
bullet('MSE - Mean Squared Error.')
bullet('RMSE - Root Mean Squared Error.')
bullet('R\u00b2 - Coefficient of Determination.')
table(
    ['Metric', 'Meaning', 'Stakeholder Interpretation'],
    [
        ['MAE', 'Average absolute prediction error', 'Typical prediction error in score units'],
        ['RMSE', 'Penalizes larger errors', 'Sensitivity to large mistakes'],
        ['R\u00b2', 'Explained variance', 'Overall explanatory performance'],
    ],
    widths=[1.2, 2.4, 2.6],
    caption_text='Table 3: Regression metrics and stakeholder interpretation',
)
callout('These are evaluation metrics to be calculated AFTER actual model execution. No values are reported here.')

page_break()

# =================================================================
# 11. HYPOTHETICAL FINDINGS
# =================================================================
heading('11. Illustrative / Hypothetical Findings', 1)
para(
    'The findings below are realistic but fictional examples used solely to demonstrate how '
    'results would be expressed and interpreted once real data is available. They are NOT '
    'measurements from any real dataset.'
)
heading('Finding 1 - Study Behavior', 3)
para(
    'Illustrative scenario: students with higher study-hour values MAY show higher predicted '
    'academic performance. This pattern would need to be validated against real data.'
)
heading('Finding 2 - Attendance', 3)
para(
    'Illustrative scenario: attendance MAY show an association with academic performance. '
    'This is a potential association, not a causal claim.'
)
heading('Finding 3 - Previous Performance', 3)
para(
    'Illustrative scenario: previous academic performance MAY be a useful predictor of future '
    'performance. This is a proposed expectation only.'
)
heading('Finding 4 - Absences', 3)
para(
    'Illustrative scenario: higher absence levels MAY be associated with lower academic outcomes. '
    'Again, this is a hypothetical observation subject to validation.'
)
heading('Finding 5 - Multiple Factors', 3)
para(
    'Student performance MAY depend on multiple interacting factors rather than any single '
    'variable alone. This motivates the use of multivariate modelling.'
)
callout('All findings above are HYPOTHETICAL and must be validated with real data before any conclusion.')

# =================================================================
# 12. MOCK-UP VISUALIZATIONS
# =================================================================
heading('12. Mock-up Visualizations', 1)
para(
    'The charts below illustrate how findings would be communicated. Every chart is explicitly '
    'marked as an illustrative example that is not based on an actual dataset.'
)
heading('Visualization 1 - Study Hours vs Final Score', 2)
figure(os.path.join(ASSETS, 'viz1_study_score.png'),
       'Figure 2: Study Hours vs Final Score (illustrative scatter plot)')
para('Purpose: show a potential relationship between study behaviour and performance.')
heading('Visualization 2 - Attendance vs Final Score', 2)
figure(os.path.join(ASSETS, 'viz2_att_score.png'),
       'Figure 3: Attendance vs Final Score (illustrative scatter plot)')
para('Purpose: explore a potential association between attendance and performance.')
heading('Visualization 3 - Final Score Distribution', 2)
figure(os.path.join(ASSETS, 'viz3_score_dist.png'),
       'Figure 4: Final Score Distribution (illustrative histogram)')
para('Purpose: understand the distribution of academic outcomes.')
heading('Visualization 4 - Performance by Student Group', 2)
figure(os.path.join(ASSETS, 'viz4_group_box.png'),
       'Figure 5: Performance by Student Group (illustrative box plot)')
para('Purpose: compare performance distributions across categories.')
heading('Visualization 5 - Feature Correlation Heatmap', 2)
figure(os.path.join(ASSETS, 'viz5_heatmap.png'),
       'Figure 6: Feature Correlation Heatmap (illustrative)')
para('Purpose: display relationships between numerical variables. The values shown are hypothetical.')
heading('Visualization 6 - Actual vs Predicted Performance', 2)
figure(os.path.join(ASSETS, 'viz6_actual_pred.png'),
       'Figure 7: Actual vs Predicted Performance (hypothetical - no model trained)')
para(
    'Purpose: demonstrate how model predictions could eventually be compared with actual values. '
    'This is a hypothetical visualization because no model has been trained.'
)

page_break()

# =================================================================
# 13. VISUALIZATION DESIGN PRINCIPLES
# =================================================================
heading('13. Visualization Design Principles', 1)
bullet('Choose charts based on the analytical question being asked.')
bullet('Avoid unnecessary decoration.')
bullet('Use clear labels and readable legends.')
bullet('Avoid misleading scales.')
bullet('Highlight the most important information.')
bullet('Include units where relevant.')
bullet('Use meaningful titles.')
bullet('Provide context around each chart.')
bullet('Add captions.')
bullet('Avoid clutter.')
para(
    'There is an important difference between visualization for analysis and visualization for '
    'executive communication. Analysis visuals are dense and exploratory, built for a data analyst '
    'to probe the data. Executive visuals are simple and focused, built for a decision-maker to '
    'grasp the key message in seconds.'
)

# =================================================================
# 14. KEY INSIGHTS
# =================================================================
heading('14. Key Insights', 1)
para(
    'Each insight is presented using a consistent framework so findings stay structured and '
    'honest. The examples below are hypothetical.'
)
for label, obs, interp, relevance, limit, action in [
    ('Insight A',
     'The illustrative scatter plot suggests higher study hours may align with higher scores.',
     'Study behaviour may be a contributing factor worth monitoring.',
     'Educators may use this to shape study-support efforts.',
     'The pattern is illustrative and not causal.',
     'Investigate with real data and confirm the direction of the relationship.'),
    ('Insight B',
     'The illustrative attendance plot suggests attendance may relate to performance.',
     'Consistent attendance may support learning continuity.',
     'Academic coordinators may consider attendance monitoring.',
     'This is an association, not a cause.',
     'Combine attendance data with other variables to study it further.'),
    ('Insight C',
     'The illustrative previous-score pattern suggests past performance may inform future outcomes.',
     'Historical performance may help identify students needing early support.',
     'Support teams may use this to prioritise attention.',
     'Past performance does not guarantee future results.',
     'Validate as a predictor within the modelling framework.'),
]:
    heading(label, 3)
    para('Observation:', bold=True); para(obs)
    para('Interpretation:', bold=True); para(interp)
    para('Business / Educational Relevance:', bold=True); para(relevance)
    para('Limitation:', bold=True); para(limit)
    para('Action:', bold=True); para(action)

page_break()

# =================================================================
# 15. ACTIONABLE INSIGHTS
# =================================================================
heading('15. Actionable Insights', 1)
para(
    'Actionable insights convert analytical observations into practical steps that stakeholders '
    'can genuinely act upon. The examples below are hypothetical.'
)
table(
    ['Insight', 'Potential Action'],
    [
        ['Attendance may be associated with academic performance.',
         'Develop attendance-monitoring and academic-support processes.'],
        ['Study behaviour may be related to performance.',
         'Provide study-planning resources and academic guidance.'],
        ['Students with weaker historical performance may require additional support.',
         'Develop early academic-support mechanisms.'],
    ],
    widths=[3.0, 3.2],
    caption_text='Table 4: Hypothetical actionable insights and potential actions',
)
callout(
    'Predictive outputs are NOT automatic decisions. All recommendations are designed to SUPPORT '
    'human judgment and decision making by educators and administrators.'
)

# =================================================================
# 16. NON-TECHNICAL STAKEHOLDER COMMUNICATION
# =================================================================
heading('16. Non-Technical Stakeholder Communication', 1)
para(
    'Data Science results must be explained to people who do not have a technical background. '
    'The key is to avoid unnecessary jargon and to describe outcomes in everyday language.'
)
para('For example, instead of saying:')
bullet('"The model achieved a lower RMSE." ')
para('Say: "The prediction system makes smaller average prediction errors."', italic=True)
para('Instead of saying:')
bullet('"Feature importance indicates\u2026"')
para('Say: "The analysis suggests which factors are most useful for making predictions."', italic=True)
table(
    ['Technical Concept', 'Plain-Language Explanation'],
    [
        ['Model', 'A prediction system trained on data'],
        ['MAE / RMSE', 'The average size of the prediction errors'],
        ['R\u00b2', 'How well the system explains differences in scores'],
        ['Cross-validation', 'Testing the system repeatedly so the result is trustworthy'],
        ['Overfitting', 'The system memorises data instead of learning a general pattern'],
        ['Data leakage', 'Using information at training time that would not be available in practice'],
    ],
    widths=[2.6, 3.6],
    caption_text='Table 5: Translating technical concepts for stakeholders',
)

# =================================================================
# 17. DATA STORYTELLING FRAMEWORK
# =================================================================
heading('17. Data Storytelling Framework', 1)
para('A narrative structure helps stakeholders remember the message.')
code(
    '1. Situation - What problem are we facing?\n'
    '       \u2193\n'
    '2. Evidence - What does the data tell us?\n'
    '       \u2193\n'
    '3. Insight - What patterns matter?\n'
    '       \u2193\n'
    '4. Meaning - Why does it matter?\n'
    '       \u2193\n'
    '5. Action - What should we do?\n'
    '       \u2193\n'
    '6. Next Step - What should happen next?'
)
caption('Figure 8: Data storytelling framework')

# =================================================================
# 18. EXECUTIVE PRESENTATION PLAN
# =================================================================
heading('18. Executive Presentation Plan', 1)
para('A proposed presentation structure that tells the project story end to end.')
numbered('Title - Student Performance Analysis & Prediction.')
numbered('Problem - Why understanding student performance matters.')
numbered('Objective - What the project aims to achieve.')
numbered('Data - What information is analysed.')
numbered('Methodology - How the analysis works.')
numbered('Key Patterns - Main illustrative findings.')
numbered('Visualization - Most important chart.')
numbered('Prediction - How ML could support prediction.')
numbered('Recommendations - What stakeholders could consider.')
numbered('Future Work - How the project could evolve.')
numbered('Conclusion - Final message.')

# =================================================================
# 19. PRESENTATION STORYBOARD
# =================================================================
heading('19. Presentation Storyboard', 1)
table(
    ['Slide', 'Message', 'Visual', 'Audience Takeaway'],
    [
        ['1', 'Project introduction', 'Title graphic', 'Understand topic'],
        ['2', 'Problem', 'Problem diagram', 'Understand need'],
        ['3', 'Objective', 'Goal diagram', 'Understand purpose'],
        ['4', 'Data', 'Data overview', 'Understand inputs'],
        ['5', 'Method', 'Workflow', 'Understand approach'],
        ['6', 'Findings', 'Charts', 'Understand patterns'],
        ['7', 'Prediction', 'ML diagram', 'Understand model'],
        ['8', 'Recommendations', 'Action framework', 'Know what to do'],
        ['9', 'Future Work', 'Roadmap', 'Understand next steps'],
        ['10', 'Conclusion', 'Summary', 'Remember key message'],
    ],
    widths=[0.8, 1.9, 1.9, 2.2],
    caption_text='Table 6: Executive presentation storyboard',
)

page_break()

# =================================================================
# 20. EXECUTIVE DASHBOARD CONCEPT
# =================================================================
heading('20. Executive Dashboard Concept', 1)
para(
    'A conceptual dashboard would give stakeholders a single view of the key indicators. '
    'Because there is no real dataset, the KPI values are marked as TBD after actual analysis.'
)
table(
    ['KPI Card', 'Status'],
    [
        ['Average Performance', 'TBD after actual analysis'],
        ['Attendance Rate', 'TBD after actual analysis'],
        ['Study Hours', 'TBD after actual analysis'],
        ['At-Risk Student Count', 'TBD after actual analysis'],
        ['Model Performance', 'TBD after actual analysis'],
    ],
    widths=[3.0, 3.2],
    caption_text='Table 7: Proposed dashboard KPI cards (values TBD)',
)
para('Proposed visual components for the dashboard:')
bullet('Score distribution.')
bullet('Attendance-performance relationship.')
bullet('Study-hour relationship.')
bullet('Student group comparison.')
bullet('Prediction overview.')
callout('No KPI values are fabricated. They remain TBD until the real analysis is performed.')

# =================================================================
# 21. RECOMMENDATIONS
# =================================================================
heading('21. Recommendations', 1)
heading('21.1 Immediate', 2)
bullet('Validate the proposed variables.')
bullet('Obtain a reliable dataset.')
bullet('Establish data-quality rules.')
bullet('Conduct the planned EDA.')
heading('21.2 Medium-Term', 2)
bullet('Train and compare multiple machine-learning models.')
bullet('Evaluate model reliability on validation and test data.')
bullet('Develop dashboards for stakeholders.')
bullet('Validate insights with domain experts.')
heading('21.3 Long-Term', 2)
bullet('Deploy the prediction system.')
bullet('Monitor model performance over time.')
bullet('Detect data drift.')
bullet('Retrain periodically.')
bullet('Integrate with educational workflows.')
callout('Recommendations support human decision making; they do not replace educators or administrators.')

page_break()

# =================================================================
# 22. FUTURE WORK
# =================================================================
heading('22. Future Work', 1)
bullet('Larger datasets.')
bullet('More historical records.')
bullet('Advanced machine-learning models.')
bullet('Explainable AI techniques.')
bullet('Personalized academic support.')
bullet('Real-time dashboards.')
bullet('Model monitoring.')
bullet('Automated retraining.')
bullet('Mobile / web integration.')
bullet('Longitudinal performance analysis.')
para(
    'These directions naturally extend the framework proposed here, moving the project from a '
    'planned methodology toward a deployed, monitored decision-support tool.'
)

# =================================================================
# 23. RISKS AND LIMITATIONS
# =================================================================
heading('23. Risks and Limitations', 1)
heading('23.1 Data Limitations', 2)
bullet('Missing information.')
bullet('Inconsistent data.')
bullet('Small sample size.')
bullet('Limited variables.')
heading('23.2 Modeling Limitations', 2)
bullet('Overfitting.')
bullet('Data leakage.')
bullet('Model bias.')
bullet('Distribution changes.')
heading('23.3 Communication Limitations', 2)
bullet('Misinterpretation of correlation.')
bullet('Overreliance on predictions.')
bullet('Oversimplification of complex findings.')
heading('23.4 Ethical Limitations', 2)
bullet('Student privacy.')
bullet('Fairness.')
bullet('Responsible use.')
bullet('Human oversight.')

# =================================================================
# 24. SUCCESS METRICS
# =================================================================
heading('24. Success Metrics', 1)
para('The future project could be evaluated across three levels.')
heading('24.1 Technical', 2)
bullet('MAE, RMSE, R\u00b2 for model accuracy.')
bullet('Cross-validation stability.')
heading('24.2 Analytical', 2)
bullet('Quality of insights.')
bullet('Data completeness.')
bullet('Reproducibility.')
heading('24.3 Business / Educational', 2)
bullet('Stakeholder usefulness.')
bullet('Interpretability.')
bullet('Actionability.')
bullet('Adoption.')
callout('No numerical targets are invented. Any specific targets would be proposed and clearly marked once the data is available.')

page_break()

# =================================================================
# 25. COMPLETE PROJECT ROADMAP
# =================================================================
heading('25. Complete Project Roadmap', 1)
code(
    'Problem\n'
    '   \u2193\n'
    'Data\n'
    '   \u2193\n'
    'Cleaning\n'
    '   \u2193\n'
    'EDA\n'
    '   \u2193\n'
    'Visualization\n'
    '   \u2193\n'
    'Feature Engineering\n'
    '   \u2193\n'
    'Machine Learning\n'
    '   \u2193\n'
    'Evaluation\n'
    '   \u2193\n'
    'Insights\n'
    '   \u2193\n'
    'Recommendations\n'
    '   \u2193\n'
    'Communication\n'
    '   \u2193\n'
    'Deployment\n'
    '   \u2193\n'
    'Monitoring'
)
caption('Figure 9: Complete Data Science project roadmap')
para(
    'This roadmap illustrates how the entire Data Science lifecycle operates as one continuous '
    'process: defining the problem, gathering and cleaning the data, exploring and visualizing '
    'patterns, building and evaluating models, generating insights and recommendations, and '
    'finally communicating, deploying, and monitoring the result.'
)

# =================================================================
# 26. CONCLUSION
# =================================================================
heading('26. Conclusion', 1)
para(
    'This Week 4 report has consolidated the Student Performance Analysis and Prediction project '
    'into a single, professional, executive-level communication document. It presents the project '
    'objective, a complete Data Science methodology, the proposed data, the exploratory analysis '
    'and visualization approach, the machine-learning framework, hypothetical findings, actionable '
    'insights, a non-technical communication strategy, a storytelling framework, presentation '
    'planning, recommendations, and future work.'
)
para(
    'Crucially, the report remains honest: because no real dataset was provided, all findings and '
    'visualizations are clearly labelled as hypothetical and illustrative. This honesty is itself '
    'a demonstration of good Data Science practice, separating real evidence from proposed '
    'concepts and preventing overstated claims.'
)
para(
    'Taken together, the earlier planning, the exploratory and visualization framework, and the '
    'machine-learning methodology converge in this document into a communication-ready package '
    'that can be shared with project managers, school administrators, academic coordinators, and '
    'mentors once real data is obtained and analysed.'
)
callout(
    'The value of a Data Science project is not only in building models, but in converting '
    'reliable data into understandable insights and responsible actions.'
)

page_break()

# =================================================================
# 27. REFERENCES
# =================================================================
heading('27. References', 1)
refs = [
    '1. Python Software Foundation. Python 3 Documentation. https://docs.python.org/3/',
    '2. Pandas Documentation. https://pandas.pydata.org/docs/',
    '3. NumPy Documentation. https://numpy.org/doc/',
    '4. Matplotlib Documentation. https://matplotlib.org/stable/contents.html',
    '5. Seaborn Documentation. https://seaborn.pydata.org/',
    '6. Plotly Documentation. https://plotly.com/python/',
    '7. scikit-learn Documentation. https://scikit-learn.org/stable/documentation.html',
    '8. Geron, A. (2022). Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow (3rd ed.). O\u2019Reilly Media.',
    '9. McKinney, W. (2022). Python for Data Analysis (3rd ed.). O\u2019Reilly Media.',
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref); run.font.size = Pt(10); run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)
para('All references are real, credible sources. No fabricated references have been included.', italic=True)

# =================================================================
# SAVE
# =================================================================
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print("Report generated successfully:", OUTPUT_FILE)

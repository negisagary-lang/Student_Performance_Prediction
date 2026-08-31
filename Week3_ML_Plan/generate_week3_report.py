"""
Week 3 - Python-Based Machine Learning Model Development and Evaluation Plan
DOCX Report Generator
Primary deliverable: Week_3_ML_Model_Development_Evaluation_Plan.docx

This is a conceptual planning document. No model has been trained or executed.
All proposed variables/outputs are labelled as illustrative.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = r"C:\Users\Lenovo\Student_Performance_Prediction\Week3_ML_Plan"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Week_3_ML_Model_Development_Evaluation_Plan.docx")

# =================================================================
# PERSONAL DETAILS (from user-provided / leave vacant)
# =================================================================
STUDENT_NAME = "Sagar Negi"
COLLEGE = "Faculty of Technology, Veer Madho Singh Bhandari"
UNIVERSITY = "Uttarakhand Technical University (UTU)"
ORG = "Yuva Internship"
PROGRAM = "Virtual Data Science Explorer Intern"
SUBMISSION_DATE = "31 August 2026"
ROLL_NUMBER = "[Roll Number]"
DEPARTMENT = "[Department]"
MENTOR = "[Mentor / Supervisor]"

doc = Document()

# ---- A4 page setup ----
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ---- Header and Footer (except first page - cover) ----
section.different_first_page_header_footer = True
header = section.header
hp = header.paragraphs[0]
hp.text = "Week 3 | Machine Learning Model Development and Evaluation Plan"
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in hp.runs:
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
run.font.size = Pt(9)
run.font.name = 'Calibri'
fld = parse_xml(
    r'<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
)
fp._p.append(fld)

# ---- Styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for level, size, color in [(1, 17, '1A3C6E'), (2, 14, '2C5F8A'), (3, 12, '344E6B')]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color)
    hs.font.bold = True


def heading(text, level):
    doc.add_heading(text, level=level)


def para(text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ' ')
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r.bold = True
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    return p


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = p.add_run(bold_prefix + ' ')
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r.bold = True
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A3C6E')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'EBF1F8')
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.italic = True
    run.bold = True
    p.paragraph_format.space_after = Pt(10)
    return p


def callout(text, color='EEF4FB', border='1A3C6E'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.italic = True
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="4" w:color="{border}"/>'
        f'</w:pBdr>'
    ))
    p.paragraph_format.space_after = Pt(8)
    return p


def centered(text, size, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def page_break():
    doc.add_page_break()


# =================================================================
# COVER PAGE
# =================================================================
for _ in range(5):
    doc.add_paragraph()

centered('WEEK 3 INTERNSHIP REPORT', 14, bold=True, space_after=20)
centered('Student Performance Analysis & Prediction', 24, bold=True, space_after=8)
centered('Python-Based Machine Learning Model Development and Evaluation Plan', 16, bold=True, space_after=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 55)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.size = Pt(10)

doc.add_paragraph()
cover = [
    ('Student Name', STUDENT_NAME),
    ('Roll Number', ROLL_NUMBER),
    ('College / Faculty', COLLEGE),
    ('University', UNIVERSITY),
    ('Department', DEPARTMENT),
    ('Internship Organization', ORG),
    ('Internship Program', PROGRAM),
    ('Mentor / Supervisor', MENTOR),
    ('Submission Date', SUBMISSION_DATE),
]
for label, value in cover:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ':  ')
    r1.font.size = Pt(12)
    r1.font.name = 'Calibri'
    r1.bold = True
    r2 = p.add_run(value)
    r2.font.size = Pt(12)
    r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()
doc.add_paragraph()
centered('This document is a planning exercise. No ML model has been trained or executed.', 10, italic=True)

page_break()

# =================================================================
# TABLE OF CONTENTS (automatic Word TOC field)
# =================================================================
heading('Table of Contents', 1)
toc_p = doc.add_paragraph()
fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>'.format(nsdecls('w')))
fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))
placeholder = parse_xml(r'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">Right-click here and choose "Update Field" to build the Table of Contents.</w:t>')
fldChar3 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
r1 = toc_p.add_run()._r
r1.append(fldChar)
r2 = toc_p.add_run()._r
r2.append(instrText)
r3 = toc_p.add_run()._r
r3.append(fldChar2)
r4 = toc_p.add_run()._r
r4.append(placeholder)
r5 = toc_p.add_run()._r
r5.append(fldChar3)

page_break()

# =================================================================
# 1. EXECUTIVE SUMMARY
# =================================================================
heading('1. Executive Summary', 1)
para(
    'This report documents the Week 3 internship task: the design of a detailed, conceptual '
    'machine-learning model development and evaluation plan for the Student Performance Analysis '
    'and Prediction project. The purpose of this week is to set out a rigorous, reproducible '
    'methodology for developing, training, validating, evaluating, and eventually deploying an ML '
    'model in Python.'
)
para(
    'The machine-learning objective is to predict a student\u2019s academic performance '
    '(proposed target: Final Score) using relevant historical and behavioral features such as '
    'study hours, attendance, previous scores, absences, and background characteristics. These '
    'features and the target are proposed and illustrative; the final set will depend on the '
    'actual dataset.'
)
para(
    'Because the target is assumed to be numerical, the primary problem formulation is '
    'supervised regression. A classification alternative is also discussed. The plan proposes:'
)
bullet('A structured preprocessing pipeline (missing values, duplicates, encoding, scaling).')
bullet('Feature engineering and feature selection strategies.')
bullet('Candidate algorithms: Linear Regression, Decision Tree, Random Forest, Gradient Boosting.')
bullet('Model training and validation using cross-validation.')
bullet('A comprehensive evaluation-metrics framework (MAE, MSE, RMSE, R\u00b2).')
bullet('A conceptual deployment, monitoring, and maintenance plan.')
callout(
    'IMPORTANT: This document presents a PROPOSED machine-learning development and evaluation '
    'plan. No model has been trained or executed as part of this planning phase. No accuracy, '
    'precision, recall, F1-score, RMSE, MAE, R\u00b2, or any other result is claimed anywhere '
    'in this document.'
)

page_break()

# =================================================================
# 2. INTRODUCTION
# =================================================================
heading('2. Introduction', 1)
para(
    'Machine Learning (ML) is a branch of artificial intelligence in which a computer system '
    'learns patterns from data without being explicitly programmed to perform a specific task. '
    'Instead of hard-coding rules, an ML model is trained on historical examples so that it can '
    'make predictions or decisions on new, unseen examples.'
)
heading('2.1 Supervised Learning', 2)
para(
    'In supervised learning, the model is trained on labelled data where both the input features '
    'and the target value are known. The model learns a mapping from inputs to outputs so that it '
    'can predict the target for new inputs.'
)
heading('2.2 Regression vs Classification', 2)
bullet('Regression: predicts a continuous numerical value (e.g., a score of 78.5).')
bullet('Classification: predicts a discrete category (e.g., \u201cHigh\u201d / \u201cMedium\u201d / \u201cLow\u201d performance).')
heading('2.3 Training, Validation, Test, and Generalization', 2)
bullet('Training: fitting the model to the training data.')
bullet('Validation: tuning model parameters and selecting between models.')
bullet('Test: final unbiased evaluation on unseen data.')
bullet('Generalization: how well the model performs on new, unseen data, beyond the data it was trained on.')
para('Machine learning fits into the broader data-science lifecycle as follows:')
code(
    'Problem Definition\n'
    '        \u2193\n'
    'Data Collection\n'
    '        \u2193\n'
    'EDA\n'
    '        \u2193\n'
    'Preprocessing\n'
    '        \u2193\n'
    'Feature Engineering\n'
    '        \u2193\n'
    'Model Selection\n'
    '        \u2193\n'
    'Model Training\n'
    '        \u2193\n'
    'Validation\n'
    '        \u2193\n'
    'Evaluation\n'
    '        \u2193\n'
    'Deployment\n'
    '        \u2193\n'
    'Monitoring'
)

page_break()

# =================================================================
# 3. PROJECT CONTEXT
# =================================================================
heading('3. Project Context', 1)
para(
    'The Student Performance Analysis and Prediction project applies data science and machine '
    'learning to educational data. Educational datasets often contain multiple variables related '
    'to student characteristics, behaviors, and outcomes. Machine learning can be used to learn '
    'patterns from historical observations and generate predictions for unseen observations.'
)
para(
    'Work completed in previous phases:'
)
bullet('Week 1: Project Planning and Strategy Development.', '')
bullet('Week 2: Exploratory Data Analysis (EDA) and Visualization Framework Design.', '')
para(
    'Week 3 builds directly on these earlier phases by proposing the machine-learning '
    'development and evaluation strategy. The EDA and visualization framework designed in '
    'Week 2 provides the foundation for understanding and cleaning the data that will feed '
    'into the model pipeline proposed here.'
)
callout(
    'Predictions from such a system are intended to support analysis and the early identification '
    'of potential academic-support needs. They are NOT intended to replace human judgment or to '
    'automatically label students.'
)

page_break()

# =================================================================
# 4. PROBLEM DEFINITION
# =================================================================
heading('4. Problem Definition', 1)
para('The proposed machine-learning problem is defined as follows:')
callout(
    'PROBLEM: Develop a machine-learning system that predicts a student\u2019s final academic '
    'performance using relevant historical and behavioral features.'
)
heading('4.1 Input', 2)
para(
    'Potential student attributes such as study hours, attendance, previous scores, assignment '
    'scores, absences, age, internet access, extracurricular activity, and family/educational '
    'background. (Proposed / illustrative.)'
)
heading('4.2 Output', 2)
para('The predicted final academic score of the student. (Proposed target: Final Score.)')
heading('4.3 ML Type', 2)
para(
    'The primary formulation is supervised regression because the Final Score is assumed to be '
    'a continuous numerical value. Regression would be preferred when the actual score is '
    'available as a continuous numerical target, because it preserves the full information in '
    'the score rather than collapsing it into coarse categories.'
)
para(
    'An alternative formulation is classification, used if the final outcome is converted into '
    'categories such as Low / Medium / High. This may be useful when approximate performance '
    'bands are sufficient for the intended decision-making.'
)

page_break()

# =================================================================
# 5. DATA REQUIREMENTS
# =================================================================
heading('5. Data Requirements', 1)
para(
    'The following table lists potential predictor features. These are proposed and illustrative; '
    'the final set will depend on the actual dataset.'
)
table(
    ['Feature', 'Type', 'Example', 'Purpose'],
    [
        ['Study Hours', 'Numerical', '5', 'Academic effort'],
        ['Attendance', 'Numerical', '90%', 'Participation'],
        ['Previous Score', 'Numerical', '78', 'Prior achievement'],
        ['Absences', 'Numerical', '4', 'Attendance behavior'],
        ['Age', 'Numerical', '20', 'Demographic context'],
        ['Internet Access', 'Categorical', 'Yes', 'Resource availability'],
        ['Extracurricular Activity', 'Categorical', 'Yes', 'Student activity'],
    ],
    widths=[2.0, 1.2, 1.0, 2.0],
)
caption('Table 1: Proposed (illustrative) input features and their purpose')
para('Target variable: Final Score (proposed / illustrative).')
callout(
    'These variables are PROPOSED only. The actual features and target will be confirmed once '
    'the real dataset is available.'
)

page_break()

# =================================================================
# 6. DATA PREPROCESSING STRATEGY
# =================================================================
heading('6. Data Preprocessing Strategy', 1)
para(
    'Preprocessing transforms raw data into a clean, consistent, model-ready format. '
    'Preprocessing is important because machine-learning algorithms make assumptions about '
    'data quality, scale, and encoding, and because errors in the raw data (missing values, '
    'duplicates, invalid records, outliers) can strongly distort model learning.'
)
heading('6.1 Preprocessing Pipeline Steps', 2)
numbered('Data loading.')
numbered('Data validation.')
numbered('Missing-value handling.')
numbered('Duplicate removal.')
numbered('Invalid-value handling.')
numbered('Outlier analysis.')
numbered('Categorical encoding.')
numbered('Numerical transformation.')
numbered('Feature scaling.')
numbered('Feature selection.')
numbered('Data-leakage prevention.')
heading('6.2 Data-Leakage Prevention in the Pipeline', 2)
para(
    'All preprocessing steps must be fit only on the training data and applied consistently to '
    'validation and test data, typically through a scikit-learn Pipeline. This prevents '
    'information from the training set leaking into evaluation data.'
)

page_break()

# =================================================================
# 7. MISSING VALUE HANDLING
# =================================================================
heading('7. Missing Value Handling', 1)
para(
    'Missing values occur when data is absent for a record. The appropriate treatment depends '
    'on the proportion of missing data, the mechanism of missingness, and the modeling goal. '
    'There is no single universally correct method.'
)
bullet('Mean imputation: replaces missing values with the column mean. Simple but sensitive to outliers.', '')
bullet('Median imputation: replaces with the column median. More robust to outliers than the mean.', '')
bullet('Mode imputation: replaces categorical missing values with the most frequent category.', '')
bullet('Advanced imputation: uses models (e.g., KNN or iterative imputation) to estimate missing values.', '')
bullet('Row removal: dropping records with missing values, only when the proportion is small and justified.', '')
callout(
    'The strategy will be selected after inspecting the actual data. Removing data or choosing '
    'an imputation method without understanding the data can bias the model.'
)

# =================================================================
# 8. CATEGORICAL ENCODING
# =================================================================
heading('8. Categorical Encoding', 1)
heading('8.1 One-Hot Encoding', 2)
para(
    'Suitable for nominal categories where there is no natural order. For example, Gender '
    '(Male / Female) becomes two binary indicator columns (male=1/0, female=1/0).'
)
heading('8.2 Ordinal Encoding', 2)
para(
    'Suitable when categories have a meaningful, known order. For example, '
    'Low < Medium < High maps to 0, 1, 2.'
)
callout(
    'Risk: Applying ordinal encoding to nominal categories incorrectly imposes an artificial '
    'order that does not exist, which can mislead distance- or gradient-based models. The correct '
    'encoding must be chosen per variable based on its true nature.'
)

# =================================================================
# 9. FEATURE SCALING
# =================================================================
heading('9. Feature Scaling', 1)
heading('9.1 Standardization', 2)
code('z = (x - mean) / standard deviation')
para('Centres the data around zero with unit variance.')
heading('9.2 Min-Max Scaling', 2)
code('x_scaled = (x - minimum) / (maximum - minimum)')
para('Rescales data to a fixed range, typically [0, 1].')
para(
    'Scaling matters for models that rely on distances (e.g., K-Nearest Neighbors, SVM) or '
    'gradient-based optimization (e.g., linear/neural models). Tree-based models (Decision Tree, '
    'Random Forest, Gradient Boosting) generally do not require standardization because they '
    'split on thresholds rather than distances.'
)

page_break()

# =================================================================
# 10. FEATURE ENGINEERING
# =================================================================
heading('10. Feature Engineering', 1)
para(
    'Feature engineering creates new, more informative variables from the raw data. Proposed '
    'examples include:'
)
bullet('Average assessment score.')
bullet('Attendance percentage.')
bullet('Absence rate.')
bullet('Study intensity.')
bullet('Performance category.')
para(
    'Feature engineering should be grounded in domain knowledge and validated against the data. '
    'Crucially, a feature must be avoided if it leaks information from the target, since this '
    'would mislead the model during training.'
)

page_break()

# =================================================================
# 11. DATA LEAKAGE
# =================================================================
heading('11. Data Leakage', 1)
para(
    'Data leakage occurs when information that would not be available at prediction time is '
    'unintentionally used during model training. This gives an optimistically biased estimate '
    'of model performance that will not hold in production.'
)
heading('11.1 Examples Relevant to Student Prediction', 2)
bullet('Using test-set statistics (e.g., mean, scaling parameters) computed from the whole dataset.')
bullet('Including a feature that already encodes the outcome (e.g., a \u201cpass/fail\u201d flag derived from the score).')
bullet('Using future-term information that would not exist at the time of prediction.')
heading('11.2 How to Prevent Leakage', 2)
bullet('Fit preprocessing (imputers, scalers, encoders) only on training data.')
bullet('Build preprocessing inside a scikit-learn Pipeline so it is applied consistently.')
bullet('Avoid using future information in feature creation.')
bullet('When feature selection is performed, do it within training folds rather than on the whole dataset.')
callout(
    'Preventing leakage is essential for obtaining an honest estimate of how the model will '
    'perform in practice.'
)

page_break()

# =================================================================
# 12. FEATURE SELECTION
# =================================================================
heading('12. Feature Selection', 1)
para(
    'Feature selection reduces the number of input variables, removing irrelevant or redundant '
    'features. This can improve model efficiency, reduce overfitting risk, and simplify '
    'interpretation.'
)
heading('12.1 Filter Methods', 2)
bullet('Correlation-based selection.')
bullet('Statistical tests.')
bullet('Variance filtering.')
heading('12.2 Wrapper Methods', 2)
bullet('Recursive Feature Elimination (RFE).')
heading('12.3 Embedded Methods', 2)
bullet('Lasso (L1) regularization.')
bullet('Tree-based feature importance.')
para(
    'The purpose is to retain informative, non-redundant variables that genuinely help predict '
    'the target, while discarding noise or duplicate information.'
)

page_break()

# =================================================================
# 13. DATASET SPLITTING
# =================================================================
heading('13. Dataset Splitting', 1)
para(
    'The dataset is divided into separate subsets used for different purposes, so that model '
    'evaluation is honest and unbiased.'
)
bullet('Training set: used to train the model.', '')
bullet('Validation set: used for model selection and hyperparameter tuning where appropriate.', '')
bullet('Test set: used for final, unbiased evaluation.', '')
para('An illustrative split could be:')
code(
    '70% Training\n'
    '15% Validation\n'
    '15% Testing'
)
para(
    'Alternatively, cross-validation may be used within the training set instead of a separate '
    'validation set. The exact split depends on dataset size and project requirements; this is '
    'illustrative only.'
)

page_break()

# =================================================================
# 14. MODEL SELECTION STRATEGY
# =================================================================
heading('14. Model Selection Strategy', 1)
para(
    'Multiple candidate algorithms will be considered so their performance can be compared on '
    'the actual data. The main candidates are described below.'
)
heading('14.1 Linear Regression', 2)
bullet('Advantages: simple, interpretable, a strong baseline.', '')
bullet('Limitations: assumes linear relationships; sensitive to certain assumptions and outliers.', '')
heading('14.2 Decision Tree', 2)
bullet('Advantages: easy to understand, captures nonlinear relationships, handles mixed feature types after preprocessing.', '')
bullet('Limitations: can overfit.', '')
heading('14.3 Random Forest', 2)
bullet('Advantages: handles nonlinear relationships, reduces single-tree variance, strong general-purpose baseline.', '')
bullet('Limitations: less interpretable than a single tree; more computationally expensive.', '')
heading('14.4 Gradient Boosting', 2)
bullet('Advantages: potential for improved predictive performance.', '')
bullet('Limitations: requires careful hyperparameter tuning.', '')
heading('14.5 Optional Classification Models', 2)
para(
    'If the target is treated as categorical, candidate classification models include: '
    'Logistic Regression, Decision Tree Classifier, and Random Forest Classifier.'
)

# =================================================================
# 15. MODEL COMPARISON TABLE
# =================================================================
heading('15. Model Comparison Table', 1)
table(
    ['Model', 'Type', 'Strengths', 'Weaknesses', 'Expected Role'],
    [
        ['Linear Regression', 'Regression', 'Simple, interpretable', 'Linear assumptions', 'Baseline'],
        ['Decision Tree', 'Regression', 'Nonlinear relationships', 'Overfitting', 'Candidate'],
        ['Random Forest', 'Regression', 'Robust, powerful', 'Less interpretable', 'Primary candidate'],
        ['Gradient Boosting', 'Regression', 'Strong performance', 'Tuning required', 'Advanced candidate'],
    ],
    widths=[1.4, 0.9, 1.6, 1.4, 1.2],
)
caption('Table 2: Proposed candidate models, strengths, weaknesses, and expected role')
callout(
    'This table does NOT claim that Random Forest actually performs best. That can only be '
    'determined by running real experiments on the actual data. Linear Regression remains an '
    'interpretable baseline and the final model will be selected only after empirical evaluation.'
)

page_break()

# =================================================================
# 16. CHOSEN MODEL STRATEGY
# =================================================================
heading('16. Chosen Model Strategy', 1)
para(
    'Based on methodological reasoning (not invented results), Random Forest Regression is '
    'proposed as a strong primary candidate. Student performance may depend on nonlinear '
    'interactions between variables such as attendance, previous performance, and study '
    'behavior, which an ensemble of trees can capture effectively.'
)
para(
    'However, Linear Regression should be retained as an interpretable baseline, and the final '
    'model should be selected only after empirical evaluation on the real dataset. This '
    'distinction is important: the proposal is a hypothesis to be tested, not a conclusion.'
)

page_break()

# =================================================================
# 17. MODEL TRAINING FRAMEWORK
# =================================================================
heading('17. Model Training Framework', 1)
numbered('Load prepared data.')
numbered('Separate features and target.')
numbered('Split the dataset.')
numbered('Build the preprocessing pipeline.')
numbered('Train a baseline model.')
numbered('Train candidate models.')
numbered('Perform cross-validation.')
numbered('Tune hyperparameters.')
numbered('Compare models.')
numbered('Select the final candidate.')
numbered('Evaluate on unseen test data.')

# =================================================================
# 18. HYPERPARAMETER TUNING
# =================================================================
heading('18. Hyperparameter Tuning', 1)
para(
    'Hyperparameters are settings that are not learned from the data but are chosen before '
    'training. Tuning finds a configuration that improves generalization.'
)
bullet('Grid Search: tests each predefined combination in a grid.', '')
bullet('Random Search: samples combinations from a defined search space.', '')
bullet('Bayesian Optimization: an advanced strategy that models performance and searches efficiently.', '')
para('For Random Forest, potential parameters to tune include:')
bullet('n_estimators: number of trees.')
bullet('max_depth: maximum tree depth.')
bullet('min_samples_split: minimum samples to split a node.')
bullet('min_samples_leaf: minimum samples in a leaf.')
bullet('max_features: number of features considered at each split.')
callout(
    'Tuning MUST be performed on validation or cross-validation data only, never on the final '
    'test set, to preserve the integrity of the final evaluation.'
)

page_break()

# =================================================================
# 19. CROSS-VALIDATION
# =================================================================
heading('19. Cross-Validation Strategy', 1)
para(
    'Cross-validation is a resampling technique that provides a more reliable estimate of '
    'model generalization than a single train/test split, by repeatedly training and validating '
    'the model on different portions of the data.'
)
heading('19.1 K-Fold Cross-Validation', 2)
code(
    'Dataset\n'
    ' \u2193\n'
    'Fold 1   Fold 2   Fold 3   Fold 4   Fold 5\n'
    '  |        |        |        |        |\n'
    '  Each fold is used as validation data once'
)
para(
    'In k-fold cross-validation, the training data is divided into k folds. The model is trained '
    'on k-1 folds and validated on the remaining fold, repeated k times so each fold is used '
    'for validation once.'
)
bullet('5-fold and 10-fold cross-validation are common choices.', '')
bullet('For classification, Stratified K-Fold preserves class proportions in each fold.', '')
bullet('Repeated cross-validation runs the process multiple times to reduce variance.', '')

# =================================================================
# 20. EVALUATION METRICS - REGRESSION
# =================================================================
heading('20. Evaluation Metrics (Regression)', 1)
heading('20.1 MAE - Mean Absolute Error', 2)
para(
    'The average absolute difference between predicted and actual values. It is easy to '
    'interpret in the target\u2019s original units.'
)
heading('20.2 MSE - Mean Squared Error', 2)
para(
    'The average of squared differences. It penalizes larger errors more strongly than MAE.'
)
heading('20.3 RMSE - Root Mean Squared Error', 2)
para(
    'The square root of MSE, returning the error in the same unit as the target. It is '
    'widely used because it is interpretable while still penalizing large errors.'
)
heading('20.4 R\u00b2 - Coefficient of Determination', 2)
para(
    'Measures the proportion of variance in the target explained by the model, relative to '
    'a baseline (the mean). Ranges typically from 0 to 1.'
)
para('Limitations of the metrics:')
bullet('MAE treats all errors equally and can hide large rare errors.')
bullet('MSE/RMSE are sensitive to outliers because errors are squared.')
bullet('R\u00b2 alone does not indicate whether the model is good in absolute terms for a given application.')
callout(
    'No metric values are reported here. These are the metrics to be computed after real '
    'model training.'
)

page_break()

# =================================================================
# 21. EVALUATION METRICS - CLASSIFICATION
# =================================================================
heading('21. Evaluation Metrics (Classification)', 1)
para(
    'If the classification alternative is adopted, the following metrics apply:'
)
bullet('Accuracy: proportion of correct predictions overall.', '')
bullet('Precision: of the positive predictions, how many were correct.', '')
bullet('Recall (Sensitivity): of the actual positives, how many were found.', '')
bullet('F1-score: the harmonic mean of precision and recall.', '')
bullet('ROC-AUC: measures the trade-off between true-positive and false-positive rates.', '')
bullet('Confusion Matrix: a table showing correct and incorrect predictions per class.', '')
callout(
    'When classes are imbalanced (e.g., far more \u201cMedium\u201d students than \u201cHigh\u201d), '
    'accuracy alone can be misleading. Metrics such as F1 and ROC-AUC should be considered.'
)

# =================================================================
# 22. METRIC SELECTION
# =================================================================
heading('22. Metric Selection', 1)
table(
    ['Problem', 'Primary Metrics', 'Supporting Metrics'],
    [
        ['Regression', 'RMSE, MAE, R\u00b2', 'MSE'],
        ['Binary Classification', 'F1, ROC-AUC', 'Accuracy, Precision, Recall'],
        ['Multiclass Classification', 'Macro-F1', 'Accuracy, Per-class metrics'],
    ],
    widths=[2.0, 1.8, 2.2],
)
caption('Table 3: Proposed primary and supporting metrics by problem type')
para(
    'Multiple metrics should be considered together because no single metric captures all '
    'aspects of model quality. A model that looks good on one metric can be poor on another.'
)

page_break()

# =================================================================
# 23. ERROR ANALYSIS
# =================================================================
heading('23. Error Analysis', 1)
para(
    'After the model is actually trained, errors will be investigated in detail. Error analysis '
    'is necessary because a model with a good average metric can still perform poorly for '
    'specific groups or specific cases.'
)
para('Proposed areas of investigation:')
bullet('Large prediction errors: which cases produce the biggest errors?')
bullet('Systematic errors: are errors consistent in one direction?')
bullet('Performance across student groups: does the model work equally well for different groups?')
bullet('Residual distribution: for regression, are residuals approximately random?')
bullet('Underprediction vs overprediction: is the model biased low or high?')
para(
    'For a regression problem, residuals (actual minus predicted) will be examined to detect '
    'patterns and systematic bias.'
)

# =================================================================
# 24. MODEL INTERPRETATION
# =================================================================
heading('24. Model Interpretation', 1)
para(
    'Interpreting the model helps explain why predictions are made and builds trust. Proposed '
    'techniques include:'
)
bullet('Feature importance: relative contribution of features.')
bullet('Permutation importance: how performance changes when a feature is shuffled.')
bullet('SHAP (SHapley Additive exPlanations): an optional advanced technique for explaining individual predictions.')
bullet('Partial dependence: an optional technique for showing how a feature affects predictions.')
callout(
    'IMPORTANT: No feature is claimed to be most important before the model is trained. '
    'Importance rankings are determined only after real training and analysis.'
)

page_break()

# =================================================================
# 25. MODEL VALIDATION
# =================================================================
heading('25. Model Validation', 1)
para(
    'A disciplined validation framework ensures the final evaluation is honest and reproducible.'
)
code(
    'Training Data\n'
    '      \u2193\n'
    'Cross-Validation\n'
    '      \u2193\n'
    'Hyperparameter Tuning\n'
    '      \u2193\n'
    'Model Selection\n'
    '      \u2193\n'
    'Final Training\n'
    '      \u2193\n'
    'Unseen Test Data\n'
    '      \u2193\n'
    'Final Evaluation'
)
callout(
    'The test set must remain untouched until the very final evaluation. It should not be used '
    'for tuning, feature selection, or model selection, otherwise the evaluation is no longer '
    'unbiased.'
)

page_break()

# =================================================================
# 26. MODEL DEPLOYMENT PLAN
# =================================================================
heading('26. Model Deployment Plan', 1)
para(
    'Although deployment is optional and not performed in this planning phase, a conceptual '
    'deployment strategy is proposed below.'
)
heading('26.1 Conceptual Architecture', 2)
code(
    'User / Application\n'
    '       \u2193\n'
    'Prediction API\n'
    '       \u2193\n'
    'Preprocessing Pipeline\n'
    '       \u2193\n'
    'Trained ML Model\n'
    '       \u2193\n'
    'Prediction\n'
    '       \u2193\n'
    'Response'
)
heading('26.2 Possible Technologies', 2)
bullet('Flask: lightweight web framework for APIs.')
bullet('FastAPI: modern, fast API framework.')
bullet('Streamlit: quick data app interface.')
bullet('Cloud deployment: hosting on cloud platforms.')
bullet('Docker: containerization for portability and reproducibility.')
callout(
    'This is a conceptual plan only. No model has been deployed.'
)

page_break()

# =================================================================
# 27. DEPLOYMENT FOR STUDENT PERFORMANCE
# =================================================================
heading('27. Deployment Concept for the Student Performance Project', 1)
para(
    'A possible future system lets a teacher or authorized user obtain a predicted performance '
    'along with supporting context, as follows:'
)
code(
    'Teacher / Authorized User\n'
    '        \u2193\n'
    'Enter Student Information\n'
    '        \u2193\n'
    'Validation\n'
    '        \u2193\n'
    'Preprocessing\n'
    '        \u2193\n'
    'ML Model\n'
    '        \u2193\n'
    'Predicted Performance\n'
    '        \u2193\n'
    'Explanation / Supporting Insights'
)
callout(
    'Predictions should support educational decision-making rather than automatically label '
    'or penalize students. The system assists educators and analysts; it does not make '
    'irreversible decisions about students automatically.'
)

page_break()

# =================================================================
# 28. MODEL MONITORING
# =================================================================
heading('28. Model Monitoring', 1)
para(
    'Once the model is deployed, its behavior and performance must be monitored over time. '
    'Relevant aspects include:'
)
bullet('Prediction performance: does accuracy hold over time?')
bullet('Data drift: does the distribution of incoming data change?')
bullet('Feature drift: do input feature distributions shift?')
bullet('Missing-value changes: does missingness change over time?')
bullet('Input distribution changes: are new records different from training data?')
bullet('Model degradation: does performance decline?')
bullet('System errors: are there failures in the pipeline?')
para(
    'Monitoring is necessary because real-world data changes over time, and a model that '
    'performed well during development can become less accurate in production.'
)

page_break()

# =================================================================
# 29. MODEL MAINTENANCE
# =================================================================
heading('29. Model Maintenance', 1)
para(
    'Model maintenance keeps the deployed model relevant and accurate. Proposed practices:'
)
bullet('Periodic retraining.')
bullet('Dataset updates.')
bullet('Version control.')
bullet('Model versioning.')
bullet('Performance tracking.')
bullet('Retraining criteria (e.g., trigger retraining when drift or performance drops beyond a threshold).')
para('Example retraining workflow:')
code(
    'New Data\n'
    '   \u2193\n'
    'Quality Check\n'
    '   \u2193\n'
    'Drift Detection\n'
    '   \u2193\n'
    'Model Evaluation\n'
    '   \u2193\n'
    'Retraining Decision\n'
    '   \u2193\n'
    'New Model Version'
)

page_break()

# =================================================================
# 30. RISK MANAGEMENT
# =================================================================
heading('30. Risk Management', 1)
para('The following risks and mitigations are proposed for the project:')
table(
    ['Risk', 'Probability', 'Impact', 'Mitigation'],
    [
        ['Missing data', 'Medium', 'High', 'Robust preprocessing'],
        ['Overfitting', 'Medium', 'High', 'Cross-validation'],
        ['Data leakage', 'Medium', 'High', 'Pipeline-based preprocessing'],
        ['Class imbalance', 'Medium', 'Medium', 'Appropriate metrics'],
        ['Small dataset', 'Medium', 'High', 'Cross-validation / data expansion'],
        ['Poor data quality', 'Medium', 'High', 'Validation rules'],
        ['Model bias', 'Medium', 'High', 'Group-level evaluation'],
        ['Concept drift', 'Low/Medium', 'Medium', 'Monitoring'],
        ['Incorrect interpretation', 'Medium', 'High', 'Explainable reporting'],
    ],
    widths=[1.8, 1.0, 0.8, 2.3],
)
caption('Table 4: Proposed risk register with mitigations')

# =================================================================
# 31. ETHICAL CONSIDERATIONS
# =================================================================
heading('31. Ethical Considerations', 1)
para(
    'A predictive system for students raises important ethical considerations that must be '
    'addressed from the start.'
)
bullet('Student privacy: protect sensitive personal information.', '')
bullet('Responsible use of predictions: predictions must not be used to harm students.', '')
bullet('Data minimization: collect and store only what is genuinely needed.', '')
bullet('Bias: be aware that historical data may contain biases.', '')
bullet('Fairness: evaluate the model across different student groups.', '')
bullet('Transparency: explain how predictions are made.', '')
bullet('Human oversight: keep humans in the decision loop.', '')
callout(
    'A predictive model should ASSIST educators and analysts; it must NOT make irreversible '
    'decisions about students automatically. Unnecessary sensitive personal data should not '
    'be included in the proposed system.'
)

page_break()

# =================================================================
# 32. REPRODUCIBILITY
# =================================================================
heading('32. Reproducibility and Documentation', 1)
para(
    'Reproducibility ensures the analysis and results can be repeated by others. The following '
    'should be documented and controlled:'
)
bullet('Python version.')
bullet('Package versions (e.g., pinned in requirements.txt).')
bullet('Random seeds for reproducibility of splits and training.')
bullet('Dataset version and source.')
bullet('Training configuration.')
bullet('Model parameters.')
bullet('Preprocessing pipeline definition.')
bullet('Experiment tracking of runs and results.')
para(
    'Reproducibility matters because it makes results verifiable, traceable, and easier to '
    'debug or extend. Without it, results cannot be trusted or rebuilt reliably.'
)

page_break()

# =================================================================
# 33. PROPOSED PROJECT STRUCTURE
# =================================================================
heading('33. Proposed Project Structure', 1)
code(
    'student-performance-ml/\n'
    '\u2502\n'
    '\u251c\u2500\u2500 data/\n'
    '\u2502   \u251c\u2500\u2500 raw/\n'
    '\u2502   \u2514\u2500\u2500 processed/\n'
    '\u2502\n'
    '\u251c\u2500\u2500 notebooks/\n'
    '\u2502   \u251c\u2500\u2500 01_eda.ipynb\n'
    '\u2502   \u2514\u2500\u2500 02_modeling.ipynb\n'
    '\u2502\n'
    '\u251c\u2500\u2500 src/\n'
    '\u2502   \u251c\u2500\u2500 preprocessing.py\n'
    '\u2502   \u251c\u2500\u2500 feature_engineering.py\n'
    '\u2502   \u251c\u2500\u2500 train.py\n'
    '\u2502   \u251c\u2500\u2500 evaluate.py\n'
    '\u2502   \u2514\u2500\u2500 predict.py\n'
    '\u2502\n'
    '\u251c\u2500\u2500 models/\n'
    '\u2502\n'
    '\u251c\u2500\u2500 reports/\n'
    '\u2502\n'
    '\u251c\u2500\u2500 requirements.txt\n'
    '\u2502\n'
    '\u251c\u2500\u2500 README.md\n'
    '\u2502\n'
    '\u2514\u2500\u2500 app.py'
)
para('Purpose of each folder:', bold=True)
bullet('data/: stores raw and processed datasets.')
bullet('notebooks/: exploratory and modeling notebooks.')
bullet('src/: reusable Python modules for preprocessing, training, evaluation, and prediction.')
bullet('models/: stores trained/serialized models.')
bullet('reports/: stores reports and visual outputs.')
bullet('requirements.txt: lists pinned package versions.')
bullet('README.md: project documentation and instructions.')
bullet('app.py: entry point for application/deployment.')

page_break()

# =================================================================
# 34. END-TO-END ML WORKFLOW DIAGRAM
# =================================================================
heading('34. End-to-End Machine Learning Workflow', 1)
para('The complete end-to-end machine-learning workflow is shown below.')
code(
    'Problem Definition\n'
    '        \u2193\n'
    'Data Collection\n'
    '        \u2193\n'
    'EDA\n'
    '        \u2193\n'
    'Data Cleaning\n'
    '        \u2193\n'
    'Feature Engineering\n'
    '        \u2193\n'
    'Train/Test Split\n'
    '        \u2193\n'
    'Preprocessing Pipeline\n'
    '        \u2193\n'
    'Baseline Model\n'
    '        \u2193\n'
    'Candidate Models\n'
    '        \u2193\n'
    'Cross-Validation\n'
    '        \u2193\n'
    'Hyperparameter Tuning\n'
    '        \u2193\n'
    'Model Comparison\n'
    '        \u2193\n'
    'Final Model\n'
    '        \u2193\n'
    'Test Evaluation\n'
    '        \u2193\n'
    'Error Analysis\n'
    '        \u2193\n'
    'Interpretation\n'
    '        \u2193\n'
    'Deployment\n'
    '        \u2193\n'
    'Monitoring\n'
    '        \u2193\n'
    'Maintenance'
)
caption('Figure 1: End-to-end machine learning workflow')

page_break()

# =================================================================
# 35. EXPECTED OUTCOMES
# =================================================================
heading('35. Expected Outcomes', 1)
para(
    'The following are PLANNED outputs. No numerical results are claimed, as the model has not '
    'been trained yet.'
)
bullet('Clean, modeling-ready dataset.')
bullet('Reproducible preprocessing pipeline.')
bullet('Multiple candidate models.')
bullet('Cross-validation results.')
bullet('Evaluation metrics (to be computed after training).')
bullet('Model comparison.')
bullet('Selected final candidate.')
bullet('Error analysis.')
bullet('Feature interpretation.')
bullet('Deployment strategy.')

# =================================================================
# 36. LIMITATIONS
# =================================================================
heading('36. Limitations', 1)
bullet('No actual dataset was available during this planning phase.')
bullet('Model performance cannot be known beforehand.')
bullet('Feature availability may change once the real dataset is obtained.')
bullet('Prediction quality depends on dataset quality.')
bullet('Historical data may contain bias.')
bullet('Correlation does not imply causation.')
bullet('Student behavior may change over time.')
bullet('Model predictions carry uncertainty.')

page_break()

# =================================================================
# 37. FUTURE SCOPE
# =================================================================
heading('37. Future Scope', 1)
bullet('Advanced ensemble models.')
bullet('Explainable AI (e.g., SHAP).')
bullet('Automated retraining.')
bullet('Real-time dashboards.')
bullet('Model monitoring systems.')
bullet('Larger datasets.')
bullet('Longitudinal student data.')
bullet('Personalized academic-support recommendations.')
para(
    'These improvements are realistic and build directly on the framework proposed in this '
    'document.'
)

# =================================================================
# 38. CONCLUSION
# =================================================================
heading('38. Conclusion', 1)
para(
    'This report has presented a comprehensive, systematic machine-learning model development '
    'and evaluation plan for the Student Performance Analysis and Prediction project. A '
    'well-structured approach to machine learning is essential: it begins with a clear problem '
    'definition, a robust preprocessing pipeline, and careful feature engineering and selection, '
    'and it progresses through model comparison, validation, and appropriate evaluation before '
    'any deployment.'
)
para(
    'The plan emphasizes the importance of proper preprocessing and data-leakage prevention, '
    'the value of comparing multiple models rather than committing to a single one, the need '
    'for disciplined validation (including the strict isolation of the test set), and the '
    'selection of metrics that match the problem type. It also stresses responsible deployment '
    'and ethical considerations, recognizing that predictions should support rather than '
    'replace human judgment.'
)
para(
    'This work connects the overall internship phases: Week 1 defined the project strategy, '
    'Week 2 provided the EDA and visualization framework for understanding the data, and '
    'Week 3 proposes the machine-learning methodology that will use those insights. The next '
    'phase will apply this plan to a real dataset, train candidate models, and evaluate them '
    'using the framework described here.'
)
para('The project follows the lifecycle:')
callout('Plan \u2192 Explore \u2192 Prepare \u2192 Model \u2192 Evaluate \u2192 Deploy \u2192 Monitor')

page_break()

# =================================================================
# 39. REFERENCES
# =================================================================
heading('39. References', 1)
refs = [
    '1. Python Official Documentation. https://docs.python.org/3/',
    '2. Pandas Documentation. https://pandas.pydata.org/docs/',
    '3. NumPy Documentation. https://numpy.org/doc/',
    '4. scikit-learn Documentation. https://scikit-learn.org/stable/documentation.html',
    '5. Matplotlib Documentation. https://matplotlib.org/stable/contents.html',
    '6. Seaborn Documentation. https://seaborn.pydata.org/',
    '7. Plotly Documentation. https://plotly.com/python/',
    '8. Geron, A. (2022). Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow (3rd ed.). O\u2019Reilly Media.',
    '9. ChatGPT (2026). I, Robot (assigned reading for ethical AI considerations).',
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

# =================================================================
# SAVE
# =================================================================
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print("Report generated successfully:", OUTPUT_FILE)

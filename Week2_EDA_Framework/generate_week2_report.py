"""
Week 2 - EDA & Visualization Framework DOCX Report Generator
Generates the primary internship deliverable:
    Week_2_EDA_Visualization_Framework.docx

No real dataset is used. All content is methodology, planning, and proposed
analysis, with any examples clearly marked as illustrative.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_BREAK

OUTPUT_DIR = r"C:\Users\Lenovo\Student_Performance_Prediction\Week2_EDA_Framework"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Week_2_EDA_Visualization_Framework.docx")

doc = Document()

# ---- Page Setup (A4) ----
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ---- Styles ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size, color in [(1, 18, '1A3C6E'), (2, 14, '2C5F8A'), (3, 12, '344E6B')]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color)
    hs.font.bold = True


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def heading(text, level):
    doc.add_heading(text, level=level)


def para(text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r.bold = True
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    return p


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r.bold = True
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    return p


def code(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A3C6E')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'EBF1F8')
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def callout(text, color='EEF4FB', border='1A3C6E'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.italic = True
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    pPr.append(shading)
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="4" w:color="{border}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    p.paragraph_format.space_after = Pt(8)
    return p


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.italic = True
    run.bold = True
    return p


def page_break():
    doc.add_page_break()


def centered(text, size, bold=False, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ================================================================
# COVER PAGE
# ================================================================
for _ in range(5):
    doc.add_paragraph()

centered('WEEK 2 INTERNSHIP TASK', 28, bold=True, space_after=6)
doc.add_paragraph()
centered('Exploratory Data Analysis (EDA)', 20, bold=True, space_after=4)
centered('& Visualization Framework Design', 20, bold=True, space_after=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 55)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.size = Pt(10)

doc.add_paragraph()
centered('Project: Student Performance Analysis & Prediction', 14, bold=True, space_after=20)

cover = [
    ('Student Name', 'Sagar Negi'),
    ('College / Faculty', 'Faculty of Technology, Veer Madho Singh Bhandari'),
    ('University', 'Uttarakhand Technical University (UTU)'),
    ('Internship Organization', 'Yuva Internship'),
    ('Internship Program', 'Virtual Data Science Explorer Intern'),
    ('Submission Date', '31 August 2026'),
]
for label, value in cover:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ':  ')
    r1.font.size = Pt(12)
    r1.font.name = 'Calibri'
    r1.bold = True
    r2 = p.add_run(value)
    r2.font.size = Pt(12)
    r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()
doc.add_paragraph()
centered('Prepared as part of the Data Science Internship Program', 11, italic=True)

page_break()

# ================================================================
# EXECUTIVE SUMMARY
# ================================================================
heading('Executive Summary', 1)

para(
    'This report documents the Week 2 internship task: the design of a comprehensive, '
    'reusable Exploratory Data Analysis (EDA) and Visualization Framework for the '
    '"Student Performance Analysis & Prediction" project. '
    'Exploratory Data Analysis is the process of investigating, summarizing, and '
    'visualizing a dataset to understand its structure, quality, and underlying patterns '
    'before any machine learning model is built.'
)

para(
    'This week\'s assignment is explicitly a framework-design task. No real dataset has '
    'been provided, so this document does not claim any actual statistics, correlations, '
    'or model results. Instead, it presents a professional methodology, the Python tools '
    'to be used, proposed analysis techniques, a visualization strategy, and a reporting '
    'and documentation framework. Any example is clearly labelled as "Illustrative" or '
    '"Proposed".'
)

para('The planned Python technology stack includes:', bold=True)
bullet('Pandas for data manipulation and summary statistics.')
bullet('NumPy for numerical and statistical operations.')
bullet('Matplotlib for static visualizations.')
bullet('Seaborn for statistical visualizations.')
bullet('Plotly for interactive visualizations.')
bullet('SciPy for statistical tests and advanced analysis.')

para('Proposed analysis techniques include:')
bullet('Univariate analysis of each variable individually.')
bullet('Bivariate analysis of relationships between two variables.')
bullet('Multivariate analysis of interactions among multiple variables.')
bullet('Missing-value analysis, duplicate detection, and outlier detection.')
bullet('Correlation analysis and visualization selection.')

callout(
    'Honesty note: Because no dataset was provided, no relationship is claimed between '
    'any variables. The framework is ready to be applied to the real Student Performance '
    'dataset in a later phase.'
)

page_break()

# ================================================================
# 1. INTRODUCTION TO EDA
# ================================================================
heading('1. Introduction to Exploratory Data Analysis', 1)

heading('1.1 What is EDA?', 2)
para(
    'Exploratory Data Analysis (EDA) is a data-driven approach to understanding a dataset '
    'through visual and statistical methods. It involves inspecting the data, assessing '
    'its quality, summarizing its main characteristics, and discovering patterns, '
    'relationships, and anomalies before formal modeling.'
)
para('The core activities of EDA include:')
bullet('Data inspection: viewing the first rows, shape, and structure.')
bullet('Data quality assessment: checking for missing values, duplicates, and errors.')
bullet('Statistical summarization: computing means, medians, spreads, and distributions.')
bullet('Pattern discovery: identifying trends and groupings.')
bullet('Relationship discovery: examining associations between variables.')
bullet('Outlier detection: finding unusual observations.')
bullet('Missing-value analysis: understanding absence of data.')
bullet('Visualization: presenting findings graphically.')

heading('1.2 Why is EDA performed BEFORE Machine Learning?', 2)
para(
    'EDA is performed before modeling because machine learning algorithms make assumptions '
    'about the data. Understanding the data first helps to ensure these assumptions hold, '
    'guides feature engineering, and prevents costly modeling mistakes.'
)
para('The conceptual flow is:')
code(
    'DATA\n'
    '  |\n'
    '  v\n'
    'UNDERSTAND\n'
    '  |\n'
    '  v\n'
    'CLEAN\n'
    '  |\n'
    '  v\n'
    'EXPLORE\n'
    '  |\n'
    '  v\n'
    'VISUALIZE\n'
    '  |\n'
    '  v\n'
    'INTERPRET\n'
    '  |\n'
    '  v\n'
    'MODEL'
)

page_break()

# ================================================================
# 2. IMPORTANCE OF EDA
# ================================================================
heading('2. Importance of EDA', 1)
para(
    'EDA is a foundational step in any data science project. The following points explain '
    'why it is essential:'
)

importance = [
    'Understanding dataset structure: knowing the number of rows, columns, and their meaning.',
    'Identifying missing values: finding gaps in data that must be addressed.',
    'Detecting duplicate records: preventing repeated rows from distorting results.',
    'Detecting outliers: identifying unusual values that may be errors or significant.',
    'Understanding distributions: knowing the shape and spread of each variable.',
    'Discovering relationships: finding how variables relate to each other.',
    'Detecting data-quality problems: catching invalid or inconsistent values.',
    'Identifying potentially useful features: spotting variables relevant to prediction.',
    'Supporting feature engineering: guiding the creation of new, informative features.',
    'Preventing incorrect modeling assumptions: avoiding models built on flawed premises.',
]
for i, point in enumerate(importance, 1):
    numbered(point)

para('The following table summarises how each EDA activity contributes to the project:', space_after=8)
table(
    ['EDA Activity', 'Purpose', 'Impact on Project'],
    [
        ['Inspect structure', 'Understand rows and columns', 'Defines feature set for modeling'],
        ['Check missing values', 'Find data gaps', 'Determines imputation strategy'],
        ['Detect duplicates', 'Find repeated records', 'Prevents biased statistics'],
        ['Detect outliers', 'Find unusual values', 'Informs cleaning and scaling'],
        ['Summarize statistics', 'Describe central tendency', 'Sets baseline for expectations'],
        ['Explore distributions', 'Understand value spread', 'Guides normalization decisions'],
        ['Analyze relationships', 'Find correlations', 'Identifies predictive features'],
        ['Visualize data', 'Present patterns visually', 'Supports insight communication'],
    ],
    widths=[1.6, 1.8, 2.6],
)
caption('Table 1: EDA activities, their purposes, and project impact')

page_break()

# ================================================================
# 3. EDA WORKFLOW
# ================================================================
heading('3. EDA Workflow', 1)
para('The EDA workflow is a structured sequence of steps. Each step is explained below.')
code(
    'Dataset Acquisition\n'
    '       |\n'
    '       v\n'
    'Data Loading\n'
    '       |\n'
    '       v\n'
    'Data Inspection\n'
    '       |\n'
    '       v\n'
    'Data Type Identification\n'
    '       |\n'
    '       v\n'
    'Data Quality Assessment\n'
    '       |\n'
    '       v\n'
    'Missing Value Analysis\n'
    '       |\n'
    '       v\n'
    'Duplicate Detection\n'
    '       |\n'
    '       v\n'
    'Outlier Detection\n'
    '       |\n'
    '       v\n'
    'Univariate Analysis\n'
    '       |\n'
    '       v\n'
    'Bivariate Analysis\n'
    '       |\n'
    '       v\n'
    'Multivariate Analysis\n'
    '       |\n'
    '       v\n'
    'Visualization\n'
    '       |\n'
    '       v\n'
    'Pattern Identification\n'
    '       |\n'
    '       v\n'
    'Insight Documentation\n'
    '       |\n'
    '       v\n'
    'EDA Report\n'
    '       |\n'
    '       v\n'
    'Machine Learning Preparation'
)
para('Stage explanations:', bold=True)
bullet('Dataset Acquisition: obtaining the data source (e.g., Kaggle).')
bullet('Data Loading: reading the data into a DataFrame using Pandas.')
bullet('Data Inspection: viewing head, shape, columns, and sample records.')
bullet('Data Type Identification: determining numeric vs categorical variables.')
bullet('Data Quality Assessment: checking for missing, duplicate, and invalid data.')
bullet('Missing Value Analysis: measuring and diagnosing missing data.')
bullet('Duplicate Detection: identifying repeated records.')
bullet('Outlier Detection: identifying extreme values.')
bullet('Univariate Analysis: studying single variables.')
bullet('Bivariate Analysis: studying two variables together.')
bullet('Multivariate Analysis: studying multiple variables together.')
bullet('Visualization: generating charts to reveal patterns.')
bullet('Pattern Identification: interpreting the visual and statistical output.')
bullet('Insight Documentation: recording findings systematically.')
bullet('EDA Report: producing a structured report of all findings.')
bullet('Machine Learning Preparation: using insights to inform modeling.')

page_break()

# ================================================================
# 4. DATA UNDERSTANDING
# ================================================================
heading('4. Data Understanding', 1)
para(
    'When first given an unfamiliar dataset, an analyst must quickly establish a baseline '
    'understanding. This includes the dataset size, column types, and immediate quality flags.'
)
para('Key questions to answer:')
numbered('How many rows does the dataset have?')
numbered('How many columns does the dataset have?')
numbered('What are the column names?')
numbered('What data types exist?')
numbered('How many unique values does each column have?')
numbered('Are there missing values?')
numbered('Are there duplicate rows?')
numbered('What do the summary statistics look like?')

para('The following Pandas commands are used to answer these questions:', space_after=8)
code(
    'import pandas as pd\n\n'
    'df = pd.read_csv("students.csv")\n\n'
    'print(df.head())          # View first 5 rows\n'
    'print(df.shape)           # Rows and columns (rows, cols)\n'
    'print(df.columns)         # Column names\n'
    'print(df.info())          # Data types and non-null counts\n'
    'print(df.describe())      # Summary statistics\n'
    'print(df.isnull().sum())  # Missing values per column\n'
    'print(df.duplicated().sum())  # Duplicate row count'
)
para('Command explanations:', bold=True)
bullet('df.head(): displays the first five rows to preview the data structure.', '')
bullet('df.shape: returns a tuple (number of rows, number of columns).', '')
bullet('df.columns: lists all column names.', '')
bullet('df.info(): shows column names, data types, and counts of non-null values.', '')
bullet('df.describe(): computes summary statistics (mean, std, min, quartiles, max).', '')
bullet('df.isnull().sum(): counts missing (NaN) values per column.', '')
bullet('df.duplicated().sum(): counts duplicate rows.', '')

page_break()

# ================================================================
# 5. DATA TYPES
# ================================================================
heading('5. Data Types', 1)
para('Correctly identifying data types determines which EDA techniques and visualizations are appropriate.')

heading('5.1 Numerical Data', 2)
bullet('Discrete: variables that take distinct, countable values (e.g., Number of Absences).')
bullet('Continuous: variables that can take any value within a range (e.g., Attendance percentage, Study Hours).')

heading('5.2 Categorical Data', 2)
bullet('Nominal: categories without a natural order (e.g., Gender, Internet Access).')
bullet('Ordinal: categories with a meaningful order (e.g., Performance Category: Low / Medium / High).')

heading('5.3 Boolean Data', 2)
para('Binary yes/no or true/false values (e.g., Extracurricular Activity: Yes/No).')

heading('5.4 Date/Time Data', 2)
para(
    'Dates and times can be parsed and used for trend analysis, periodicity detection, '
    'and seasonal patterns (e.g., exam dates, enrollment years).'
)

heading('5.5 Text Data', 2)
para(
    'Free-text variables (e.g., comments, feedback) require separate preprocessing such as '
    'tokenization and natural language processing before they can be used in EDA.'
)

para('The table below links data types to suitable analysis techniques:', space_after=8)
table(
    ['Data Type', 'Example', 'Suitable EDA Techniques', 'Suitable Visualizations'],
    [
        ['Discrete', 'Absences', 'Frequency counts, describe()', 'Bar chart, histogram'],
        ['Continuous', 'Study Hours', 'Summary stats, histograms', 'Histogram, box plot, KDE'],
        ['Nominal', 'Gender', 'Value counts, crosstabs', 'Bar chart, pie chart'],
        ['Ordinal', 'Performance Category', 'Ranked counts', 'Ordered bar chart'],
        ['Boolean', 'Internet Access', 'Yes/No counts', 'Bar chart, stacked chart'],
        ['Date/Time', 'Enrollment Date', 'Time aggregation', 'Line chart, trend plot'],
        ['Text', 'Feedback', 'Tokenize, word frequency', 'Word cloud, bar chart'],
    ],
    widths=[1.2, 1.4, 1.8, 1.8],
)
caption('Table 2: Data types, examples, and suitable EDA techniques')

page_break()

# ================================================================
# 6. DATA QUALITY ASSESSMENT
# ================================================================
heading('6. Data Quality Assessment', 1)
para(
    'Data quality directly affects the reliability of analysis and modeling. A systematic '
    'framework assesses missing values, duplicates, invalid values, and outliers.'
)

heading('6.1 Missing Values', 2)
para('Missing values occur when data is absent for a variable. The Pandas code to detect them is:')
code('df.isnull().sum()')
para('Missing-data mechanisms (high-level):')
bullet('MCAR (Missing Completely At Random): absence is unrelated to any variable.')
bullet('MAR (Missing At Random): absence depends on observed variables.')
bullet('MNAR (Missing Not At Random): absence depends on the missing value itself.')
para('Possible treatment strategies:')
bullet('Remove records: delete rows with missing values (if few).')
bullet('Mean imputation: replace with the column mean.')
bullet('Median imputation: replace with the column median (robust to outliers).')
bullet('Mode imputation: replace categorical values with the most frequent.')
bullet('Forward/backward fill: propagate surrounding values (for time series).')
bullet('Advanced imputation: use models (e.g., KNN or MICE) to estimate missing values.')
callout(
    'The chosen strategy depends on the dataset, the proportion of missing data, and the '
    'modeling goal. No single method fits every situation.'
)

heading('6.2 Duplicate Data', 2)
para('Duplicate records are identical (or near-identical) rows. Detection and removal:')
code(
    'print(df.duplicated().sum())   # count duplicates\n'
    'df = df.drop_duplicates()      # remove duplicates'
)
para(
    'Duplicate records can distort means, counts, correlations, visualizations, and model '
    'training because repeated rows inflate the influence of the duplicated observation.'
)

heading('6.3 Invalid Data', 2)
para('Invalid data violates domain rules. Examples:')
bullet('Age = -10 (impossible negative age).')
bullet('Attendance = 150% (exceeds 100%).')
bullet('Score = 500 / 100 (exceeds maximum).')
bullet('Negative study hours (logically impossible).')
para('Enforcement approaches:')
bullet('Validation rules: define allowed value ranges.')
bullet('Range checks: reject values outside a valid range.')
bullet('Logical consistency checks: ensure related fields are coherent.')
bullet('Domain validation: apply subject-matter expertise.')

page_break()

# ================================================================
# 7. OUTLIER DETECTION
# ================================================================
heading('7. Outlier Detection', 1)
para(
    'Outliers are observations that differ significantly from the rest of the data. '
    'They can be detected using several methods.'
)

heading('7.1 IQR Method', 2)
para('The Interquartile Range (IQR) method identifies outliers relative to quartiles:')
code(
    'Q1 = 25th percentile\n'
    'Q3 = 75th percentile\n'
    'IQR = Q3 - Q1\n'
    'Lower Bound = Q1 - 1.5 * IQR\n'
    'Upper Bound = Q3 + 1.5 * IQR\n\n'
    '# Any value outside [Lower Bound, Upper Bound] is an outlier'
)

heading('7.2 Other Methods', 2)
bullet('Box plots: visually show spread and flag extreme values.')
bullet('Z-score: counts standard deviations from the mean (e.g., |z| > 3 as outlier).')
bullet('Percentiles: values beyond e.g. the 1st and 99th percentiles.')
bullet('Domain-based thresholds: values outside physically/logically possible ranges.')

callout(
    'IMPORTANT: Outliers should NOT automatically be deleted. An outlier may be a '
    'data-entry error, a rare but valid observation, or a meaningful exceptional case. '
    'Investigate before deciding to remove, cap, or keep it.',
    color='FFF7E6',
    border='ED7D31'
)

page_break()

# ================================================================
# 8. UNIVARIATE ANALYSIS
# ================================================================
heading('8. Univariate Analysis', 1)
para('Univariate analysis studies one variable at a time.')

heading('8.1 Numerical Variables', 2)
para('Key statistics calculated:')
bullet('Mean, Median, Mode: measures of central tendency.')
bullet('Minimum, Maximum, Range: measure of extent.')
bullet('Variance, Standard Deviation: measures of spread.')
bullet('Quartiles and IQR: measure of spread and position.')
code(
    'df["Final_Score"].describe()'
)

heading('8.2 Categorical Variables', 2)
para('Key statistics calculated:')
bullet('Frequency: count of each category.')
bullet('Percentage: proportion of each category.')
bullet('Unique categories: number of distinct values.')
code(
    'df["Gender"].value_counts()\n'
    'df["Gender"].value_counts(normalize=True)'
)

# ================================================================
# 9. UNIVARIATE VISUALIZATION
# ================================================================
heading('9. Univariate Visualization', 1)
para('Appropriate visualizations reveal the distribution of a single variable.')

heading('9.1 Histogram', 2)
para('Purpose: understand the frequency distribution of a numerical variable.')
para('Example: Distribution of Final Score')
code(
    "import matplotlib.pyplot as plt\n"
    "import seaborn as sns\n\n"
    "sns.histplot(df['Final_Score'], kde=True, bins=20)\n"
    "plt.title('Distribution of Final Score')\n"
    "plt.show()"
)

heading('9.2 Box Plot', 2)
para('Purpose: identify spread, quartiles, and potential outliers.')
para('Example: Box plot of Study Hours')
code(
    "sns.boxplot(y=df['Study_Hours'])\n"
    "plt.title('Box Plot of Study Hours')\n"
    "plt.show()"
)

heading('9.3 Bar Chart', 2)
para('Purpose: compare frequencies of categorical values.')
para('Example: Gender distribution')
code(
    "sns.barplot(x=df['Gender'].value_counts().index,\n"
    "            y=df['Gender'].value_counts().values)\n"
    "plt.show()"
)

heading('9.4 KDE Plot', 2)
para('Purpose: understand the smooth shape of a distribution.')
para('Example: KDE of Attendance')
code(
    "sns.kdeplot(df['Attendance'], fill=True)\n"
    "plt.title('KDE of Attendance')\n"
    "plt.show()"
)

page_break()

# ================================================================
# 10. BIVARIATE ANALYSIS
# ================================================================
heading('10. Bivariate Analysis', 1)
para('Bivariate analysis studies the relationship between two variables.')

heading('10.1 Numerical + Numerical', 2)
para('Examples: Study Hours vs Final Score; Attendance vs Final Score; Previous Score vs Final Score.')
para('Techniques: Scatter plot, correlation, regression line.')
code(
    "sns.scatterplot(x=df['Study_Hours'], y=df['Final_Score'])\n"
    "plt.title('Study Hours vs Final Score')\n"
    "plt.show()\n\n"
    "# With a regression line\n"
    "sns.regplot(x=df['Study_Hours'], y=df['Final_Score'])\n"
    "plt.show()"
)

heading('10.2 Categorical + Numerical', 2)
para('Examples: Gender vs Final Score; Internet Access vs Final Score.')
para('Techniques: Box plot, violin plot, bar chart.')
code(
    "sns.boxplot(x=df['Gender'], y=df['Final_Score'])\n"
    "plt.show()\n\n"
    "sns.violinplot(x=df['Internet_Access'], y=df['Final_Score'])\n"
    "plt.show()"
)

heading('10.3 Categorical + Categorical', 2)
para('Example: Gender vs Extracurricular Activity.')
para('Techniques: Grouped bar chart, stacked bar chart, crosstab.')
code(
    "import pandas as pd\n"
    "ct = pd.crosstab(df['Gender'], df['Internet_Access'])\n"
    "ct.plot(kind='bar', stacked=True)\n"
    "plt.show()"
)

page_break()

# ================================================================
# 11. MULTIVARIATE ANALYSIS
# ================================================================
heading('11. Multivariate Analysis', 1)
para(
    'Multivariate analysis studies interactions among multiple variables simultaneously. '
    'Example: Study Hours + Attendance + Previous Score together influence Final Score.'
)
para('Techniques:')
bullet('Correlation matrix: measures pairwise relationships among many variables.')
bullet('Heatmap: visualizes the correlation matrix compactly.')
bullet('Pair plot: grid of scatter plots for all pairwise combinations.')
bullet('Grouped analysis: compare a variable across multiple categories.')
bullet('Faceted plots: visualize sub-groups side by side.')
bullet('Feature interaction analysis: examine how combinations of features relate.')
code(
    "import seaborn as sns\n\n"
    "num_cols = df.select_dtypes(include='number').columns\n"
    "corr = df[num_cols].corr()\n"
    "sns.heatmap(corr, annot=True, cmap='coolwarm')\n"
    "plt.title('Correlation Heatmap')\n"
    "plt.show()\n\n"
    "# Pair plot of numerical variables\n"
    "sns.pairplot(df[num_cols])"
)
para(
    'Multivariate analysis is especially useful before machine learning because it reveals '
    'multicollinearity among predictors and helps identify which combinations of features '
    'best explain the target variable.'
)

page_break()

# ================================================================
# 12. CORRELATION ANALYSIS
# ================================================================
heading('12. Correlation Analysis', 1)
para('Correlation measures the strength and direction of a linear relationship between two variables.')
bullet('Pearson correlation: measures linear relationships.')
bullet('Spearman correlation: measures monotonic (rank-based) relationships.')
para('The correlation coefficient ranges from -1 to +1:')
bullet('+1: perfect positive correlation (both increase together).')
bullet('0: no linear correlation.')
bullet('-1: perfect negative correlation (one increases, other decreases).')
para('Strength classification:')
bullet('|r| near 1: strong correlation.')
bullet('|r| near 0: weak correlation.')
callout(
    'CRITICAL: Correlation does NOT imply causation. Two correlated variables do not '
    'necessarily have a cause-and-effect relationship. This principle must be respected '
    'when interpreting the real data.',
    color='FFF7E6',
    border='ED7D31'
)
para(
    'Illustrative example: With a real dataset, if Attendance and Final Score show a high '
    'positive Pearson correlation, one might observe that higher attendance tends to '
    'coincide with higher scores. However, this observation alone does not prove that '
    'attendance causes better scores - other factors may be involved. (This is a proposed '
    'interpretation, not a claim made from real data.)'
)

page_break()

# ================================================================
# 13. VISUALIZATION STRATEGY
# ================================================================
heading('13. Visualization Strategy', 1)
para('Choosing the right chart depends on the analytical question being asked.')

table(
    ['Analytical Question', 'Recommended Plot', 'Reason'],
    [
        ['What is the distribution?', 'Histogram', 'Shows frequency distribution'],
        ['Are there outliers?', 'Box Plot', 'Shows spread and extreme values'],
        ['Are two numerical variables related?', 'Scatter Plot', 'Shows relationship'],
        ['How do categories compare?', 'Bar Chart', 'Easy comparison'],
        ['How do distributions differ?', 'Violin Plot', 'Shows distribution shape'],
        ['What variables are correlated?', 'Heatmap', 'Compact correlation overview'],
        ['How does something change over time?', 'Line Chart', 'Shows trends'],
        ['How do multiple variables interact?', 'Pair Plot', 'Multi-variable exploration'],
    ],
    widths=[2.0, 1.5, 2.5],
)
caption('Table 3: Visualization selection framework')

page_break()

# ================================================================
# 14-16. LIBRARY FRAMEWORKS
# ================================================================
heading('14. Python Visualization Libraries', 1)
para(
    'The framework relies on several complementary Python libraries. Their purposes are '
    'summarised below.'
)
table(
    ['Library', 'Purpose', 'Example Usage'],
    [
        ['Pandas', 'Data manipulation', 'Loading and cleaning data'],
        ['NumPy', 'Numerical operations', 'Statistical calculations'],
        ['Matplotlib', 'Basic visualization', 'Histograms, line charts'],
        ['Seaborn', 'Statistical visualization', 'Heatmaps, boxplots'],
        ['Plotly', 'Interactive visualization', 'Interactive dashboards'],
        ['SciPy', 'Statistical analysis', 'Correlation and statistical tests'],
    ],
    widths=[1.5, 2.2, 2.3],
)
caption('Table 4: Core libraries and their purposes')

heading('14.1 Matplotlib Framework', 2)
para(
    'Matplotlib is the foundational plotting library. It provides control over figures, '
    'axes, titles, labels, legends, and gridlines, and allows figures to be saved to file.'
)
para('Key concepts:')
bullet('Figure: the top-level container for all plot elements.')
bullet('Axes: the region where data is drawn.')
bullet('Title, labels, legend: textual annotations.')
bullet('Grid: guides for reading values.')
bullet('Figure size: controls dimensions via figsize.')
bullet('Saving: plt.savefig() exports the figure.')
code(
    "import matplotlib.pyplot as plt\n\n"
    "plt.figure(figsize=(8, 5))\n"
    "plt.hist(df['Final_Score'])\n"
    "plt.title('Distribution of Final Scores')\n"
    "plt.xlabel('Final Score')\n"
    "plt.ylabel('Frequency')\n"
    "plt.grid(True)\n"
    "plt.savefig('final_score_dist.png')\n"
    "plt.show()"
)

heading('14.2 Seaborn Framework', 2)
para(
    'Seaborn builds on Matplotlib to provide statistical visualizations with better '
    'default aesthetics. Key functions include:'
)
bullet('sns.histplot(): histogram with optional KDE.')
bullet('sns.boxplot(): box plot for spread and outliers.')
bullet('sns.scatterplot(): scatter plot for relationships.')
bullet('sns.heatmap(): correlation matrix visualization.')
bullet('sns.pairplot(): grid of pairwise plots.')
bullet('sns.barplot(): bar chart for category comparisons.')

heading('14.3 Plotly Framework', 2)
para(
    'Plotly provides interactive visualizations. Its advantages include hover information, '
    'zooming, filtering, and interactive exploration of data.'
)
code(
    "import plotly.express as px\n\n"
    "# Interactive scatter plot\n"
    "fig = px.scatter(df, x='Study_Hours', y='Final_Score')\n"
    "fig.show()\n\n"
    "# Interactive bar chart / histogram\n"
    "fig2 = px.histogram(df, x='Final_Score')\n"
    "fig2.show()"
)

page_break()

# ================================================================
# 15. PROPOSED VISUALIZATIONS
# ================================================================
heading('15. Proposed Visualizations for Student Performance', 1)
para(
    'The following visualizations are proposed for when the real dataset becomes available. '
    'No actual relationship is claimed until real data is analyzed.'
)

table(
    ['#', 'Proposed Visualization', 'Analysis Type', 'Purpose'],
    [
        ['1', 'Final Score Distribution', 'Univariate', 'Understand score spread'],
        ['2', 'Study Hours vs Final Score', 'Bivariate', 'Explore study-time association'],
        ['3', 'Attendance vs Final Score', 'Bivariate', 'Explore attendance association'],
        ['4', 'Previous Score vs Final Score', 'Bivariate', 'Explore prior-score link'],
        ['5', 'Average Score by Gender', 'Bivariate', 'Compare across genders'],
        ['6', 'Score Distribution by Performance Category', 'Bivariate', 'Compare performance groups'],
        ['7', 'Absences vs Final Score', 'Bivariate', 'Explore absences link'],
        ['8', 'Correlation Heatmap', 'Multivariate', 'Overview of correlations'],
        ['9', 'Feature Importance', 'Modeling phase', 'Identify predictive features'],
    ],
    widths=[0.4, 2.4, 1.3, 2.0],
)
caption('Table 5: Proposed visualizations')
callout(
    'All the above visualizations are PROPOSED. They will only be generated and interpreted '
    'after the real Student Performance dataset is supplied.'
)

page_break()

# ================================================================
# 16. EDA QUESTIONS FRAMEWORK
# ================================================================
heading('16. EDA Questions Framework', 1)
para('A structured set of questions guides the EDA process.')

heading('16.1 Dataset-Level Questions', 2)
bullet('How large is the dataset?')
bullet('What variables are available?')
bullet('Which variables are numerical?')
bullet('Which are categorical?')
bullet('Are there missing values?')
bullet('Are there duplicates?')

heading('16.2 Numerical Questions', 2)
bullet('What is the distribution?')
bullet('What is the average?')
bullet('What is the spread?')
bullet('Are there outliers?')

heading('16.3 Relationship Questions', 2)
bullet('Are study hours associated with scores?')
bullet('Is attendance associated with performance?')
bullet('Are previous scores related to final scores?')
bullet('Are there strong correlations between predictors?')

heading('16.4 Data Quality Questions', 2)
bullet('Are values valid?')
bullet('Are there inconsistent categories?')
bullet('Are units consistent?')

page_break()

# ================================================================
# 17. REUSABLE EDA CODE FRAMEWORK
# ================================================================
heading('17. Reusable EDA Python Code Framework', 1)
para(
    'The framework is implemented as reusable functions, which improve efficiency and '
    'consistency across analyses. The reusable module is provided as eda_framework.py.'
)
para('The core reusable functions include:')
code(
    "def dataset_summary(df):\n"
    "    pass\n\n"
    "def missing_value_report(df):\n"
    "    pass\n\n"
    "def numerical_summary(df):\n"
    "    pass\n\n"
    "def categorical_summary(df):\n"
    "    pass\n\n"
    "def detect_outliers(df):\n"
    "    pass\n\n"
    "def correlation_matrix(df):\n"
    "    pass\n\n"
    "def run_full_eda(df):\n"
    "    pass"
)
para(
    'Advantages of reusable functions: they standardize the analysis steps, reduce '
    'repetitive code, minimize errors, and can be applied to any dataset simply by '
    'calling them. This makes the EDA reproducible and maintainable.'
)

heading('17.1 Standard EDA Template Structure', 2)
para('A reusable analysis template follows this structure:')
numbered('Import libraries.')
numbered('Load dataset.')
numbered('Inspect dataset.')
numbered('Data quality checks.')
numbered('Missing-value analysis.')
numbered('Duplicate analysis.')
numbered('Statistical summary.')
numbered('Univariate analysis.')
numbered('Bivariate analysis.')
numbered('Multivariate analysis.')
numbered('Visualization.')
numbered('Insight documentation.')

page_break()

# ================================================================
# 18. REPORTING FRAMEWORK
# ================================================================
heading('18. Reporting Framework', 1)
para(
    'EDA findings must be communicated through a structured report. The proposed reporting '
    'structure is:'
)
numbered('Executive Summary.')
numbered('Dataset Overview.')
numbered('Data Quality Assessment.')
numbered('Data Cleaning Summary.')
numbered('Univariate Analysis.')
numbered('Bivariate Analysis.')
numbered('Multivariate Analysis.')
numbered('Visualization Findings.')
numbered('Key Insights.')
numbered('Limitations.')
numbered('Recommendations.')
numbered('Conclusion.')

para('For each important chart, findings are documented using:')
code(
    "Chart -> Observation -> Interpretation -> Implication"
)
bullet('Chart: what chart was used.')
bullet('Observation: what is visually present.')
bullet('Interpretation: what it means statistically.')
bullet('Implication: how it affects modeling or decisions.')

page_break()

# ================================================================
# 19. DOCUMENTATION STANDARDS
# ================================================================
heading('19. Documentation Standards', 1)
para('Proper documentation ensures that EDA results are reproducible and auditable.')
para('Each analysis should record:')
bullet('Analysis objective.')
bullet('Variable(s) involved.')
bullet('Method used.')
bullet('Visualization used.')
bullet('Observation.')
bullet('Interpretation.')
bullet('Limitation.')
bullet('Next action.')

para('To maintain reproducibility, document:')
bullet('Python version used.')
bullet('Library versions used.')
bullet('Dataset source.')
bullet('Dataset version/date.')
bullet('Cleaning decisions made.')
bullet('Analysis decisions made.')
para(
    'These records allow the analysis to be repeated by others and ensure that any '
    'conclusions rest on clearly stated, reproducible steps.'
)

page_break()

# ================================================================
# 20. HTML FORMATTING REQUIREMENT
# ================================================================
heading('20. HTML Formatting for Data-Analysis Reporting', 1)
para(
    'HTML formatting is a valuable tool for data-analysis reporting. The assignment '
    'specifically evaluates the use of HTML formatting, so this section explains how HTML '
    'improves reports and demonstrates the required elements.'
)
para('Common HTML elements used in reports:')
code(
    "<h1>Main Heading</h1>\n"
    "<h2>Section Heading</h2>\n"
    "<p>A paragraph of text.</p>\n"
    "<table>\n"
    "  <tr><th>Header</th><th>Header</th></tr>\n"
    "  <tr><td>Data</td><td>Data</td></tr>\n"
    "</table>\n"
    "<ul><li>Bullet item</li></ul>\n"
    "<ol><li>Numbered item</li></ol>\n"
    "<strong>Bold / important text</strong>\n"
    "<em>Italic / emphasized text</em>\n"
    "<img src='chart.png' alt='Chart'>"
)
para('How HTML integrates with the data-analysis workflow:')
bullet('Jupyter Notebook: Markdown cells support HTML for rich formatting.')
bullet('Markdown: HTML can be embedded for tables and styling.')
bullet('HTML reports: data analysis can be exported as standalone HTML documents.')
bullet('Interactive Plotly: Plotly figures export as interactive HTML with hover and zoom.')
para(
    'A full, valid standalone HTML demonstration file (eda_report_template.html) is '
    'provided alongside this document.'
)

page_break()

# ================================================================
# 21. CONCLUSION
# ================================================================
heading('21. Conclusion', 1)
para(
    'This report has presented a comprehensive, professional Exploratory Data Analysis '
    'and Visualization Framework designed for the "Student Performance Analysis & '
    'Prediction" project. The framework covers the definition and importance of EDA, '
    'the complete EDA workflow, data understanding, data types, data quality assessment, '
    'outlier detection, univariate/bivariate/multivariate analysis, correlation analysis, '
    'and a structured visualization strategy.'
)
para(
    'A reusable Python code framework (eda_framework.py) and a sample analysis template '
    '(sample_eda_template.ipynb) have been developed to operationalize the methodology. '
    'The reporting and documentation standards ensure that all findings are reproducible '
    'and clearly communicated, and the HTML formatting discussion demonstrates how '
    'professional, interactive reports can be produced.'
)
para(
    'Importantly, this week\'s task is a framework-design assignment. No real dataset was '
    'provided, and therefore no actual statistics or relationships have been claimed. '
    'All examples are clearly labelled as illustrative or proposed. The framework is '
    'fully prepared to be applied to the real Student Performance dataset in the next '
    'phase of the internship, at which point genuine insights, cleaning decisions, and '
    'modeling steps can be documented.'
)

doc.add_paragraph()
callout(
    'Deliverables accompanying this report: eda_framework.py, sample_eda_template.ipynb, '
    'eda_report_template.html, and README.md.'
)

page_break()

# ================================================================
# 22. REFERENCES
# ================================================================
heading('22. References', 1)
refs = [
    '1. Pandas Documentation. https://pandas.pydata.org/docs/',
    '2. NumPy Documentation. https://numpy.org/doc/',
    '3. Matplotlib Documentation. https://matplotlib.org/stable/contents.html',
    '4. Seaborn Documentation. https://seaborn.pydata.org/',
    '5. Plotly Documentation. https://plotly.com/python/',
    '6. SciPy Documentation. https://docs.scipy.org/doc/scipy/',
    '7. Python Official Documentation. https://docs.python.org/3/',
    '8. Jupyter Documentation. https://jupyter.org/documentation',
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

# ================================================================
# SAVE
# ================================================================
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print("Report generated successfully:", OUTPUT_FILE)
